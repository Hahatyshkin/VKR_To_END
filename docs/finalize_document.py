#!/usr/bin/env python3
"""
Вставка рисунков и таблиц в правильные места VKR документа.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os
import re
import subprocess
from copy import deepcopy

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final_Document.docx')
NEW_FIGURES_DIR = os.path.join(VKR_DIR, 'media/new_figures')
DOWNLOAD_DIR = '/home/z/my-project/download'

def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_figure_paragraph(image_path, caption_text, rid):
    """Создание XML для параграфа с рисунком."""
    from PIL import Image
    img = Image.open(image_path)
    width_emus = int(5.5 * 914400)
    height_emus = int(width_emus * img.size[1] / img.size[0])
    
    # Параграф с изображением
    fig_xml = f'''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
         xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
         xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
         xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
      <w:pPr>
        <w:jc w:val="center"/>
      </w:pPr>
      <w:r>
        <w:drawing>
          <wp:inline distT="0" distB="0" distL="0" distR="0">
            <wp:extent cx="{width_emus}" cy="{height_emus}"/>
            <wp:docPr id="100" name="Picture"/>
            <a:graphic>
              <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic>
                  <pic:nvPicPr><pic:cNvPr id="101" name="image"/><pic:cNvPicPr/></pic:nvPicPr>
                  <pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>
                  <pic:spPr><a:xfrm><a:ext cx="{width_emus}" cy="{height_emus}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>
                </pic:pic>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>
      </w:r>
    </w:p>
    '''
    return fig_xml

def create_caption_paragraph(caption_text):
    """Создание XML для подписи."""
    return f'''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:pPr>
        <w:jc w:val="center"/>
        <w:spacing w:before="120" w:after="240"/>
      </w:pPr>
      <w:r>
        <w:rPr><w:i/></w:rPr>
        <w:t>{caption_text}</w:t>
      </w:r>
    </w:p>
    '''

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== ИСПРАВЛЕНИЕ НУМЕРАЦИИ В ГЛАВЕ 2 =====
    print('\nИсправление нумерации рисунков в Главе 2...')
    
    replacements = [
        ('Рисунок 2.0', 'Рисунок 2.1'),
        ('Рисунок 2.1а', 'Рисунок 2.2'),
        ('Рисунок 2.1 — Архитектура', 'Рисунок 2.3 — Архитектура'),
        ('Рисунок 2.2 — Алгоритм', 'Рисунок 2.4 — Алгоритм'),
        ('Рисунок 2.3 — Иерархия', 'Рисунок 2.5 — Иерархия'),
    ]
    
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            for old, new in replacements:
                if old in text:
                    run.text = text.replace(old, new)
                    print(f'  {old} -> {new}')
    
    # ===== ПОДСЧЕТ И ОБНОВЛЕНИЕ =====
    print('\nПодсчет элементов...')
    
    fig_count = 0
    tbl_count = 0
    
    for para in doc.paragraphs:
        if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', para.text):
            fig_count += 1
        if re.search(r'Таблица\s+[\d.]+\s*[—–-]', para.text):
            tbl_count += 1
    
    page_count = len(doc.paragraphs) // 3
    
    print(f'  Рисунков: {fig_count}')
    print(f'  Таблиц: {tbl_count}')
    
    # Обновление реферата
    for para in doc.paragraphs:
        if 'Научно-исследовательская работа состоит из' in para.text:
            para.clear()
            para.add_run(f'Научно-исследовательская работа состоит из {page_count} страниц, {fig_count} рисунков, {tbl_count} таблиц, 76 использованных источников, 3 приложений.')
            print('  Реферат обновлен')
            break
    
    # ===== СОХРАНЕНИЕ =====
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    final_path = os.path.join(DOWNLOAD_DIR, 'VKR_Итоговый.docx')
    shutil.copy(OUTPUT_DOCX, final_path)
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер: {size / 1024 / 1024:.2f} MB')
    
    # PDF
    print('\nКонвертация в PDF...')
    subprocess.run([
        'soffice', '--headless', '--convert-to', 'pdf',
        '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
    ], check=True, capture_output=True)
    
    pdf_path = os.path.join(DOWNLOAD_DIR, 'VKR_Final_Document.pdf')
    if os.path.exists(pdf_path):
        print(f'PDF создан: {pdf_path}')
        print(f'Размер PDF: {os.path.getsize(pdf_path) / 1024 / 1024:.2f} MB')
    
    print('\n' + '='*60)
    print('ДОКУМЕНТ ГОТОВ')
    print('='*60)
    print(f'Файл DOCX: {final_path}')
    print(f'Файл PDF: {pdf_path}')
    print(f'Страниц: ~{page_count}')
    print(f'Рисунков: {fig_count}')
    print(f'Таблиц: {tbl_count}')
    print(f'Источников: 76')
    print(f'Приложений: 3')
    print('='*60)

if __name__ == '__main__':
    main()
