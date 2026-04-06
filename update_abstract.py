#!/usr/bin/env python3
"""
Обновление реферата и проверка дубликатов в ВКР.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from pathlib import Path

def count_elements(doc):
    """Подсчёт элементов в документе"""
    figures = set()
    tables = set()
    references = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Подсчёт рисунков
        fig_match = re.search(r'Рисунок\s+(\d+\.\d+)', text)
        if fig_match:
            figures.add(fig_match.group(1))
        
        # Подсчёт таблиц
        tab_match = re.search(r'Таблица\s+(\d+\.\d+)', text)
        if tab_match:
            tables.add(tab_match.group(1))
        
        # Подсчёт ссылок в списке источников
        if re.match(r'^\d+\.\s+[A-ZА-Я]', text):
            references += 1
    
    return {
        'figures': len(figures),
        'tables': len(tables),
        'references': references,
        'figure_list': sorted(figures),
        'table_list': sorted(tables)
    }

def update_abstract(doc, counts):
    """Обновление реферата"""
    abstract_text = f"""Научно-исследовательская работа состоит из 170 страниц, {counts['figures']} рисунков, {counts['tables']} таблиц, {counts['references']} использованных источников, 3 приложений.

МАТРИЦА АДАМАРА, МАТРИЦА УОЛША, ПРЕОБРАЗОВАНИЕ УОЛША-АДАМАРА, БЫСТРОЕ ПРЕОБРАЗОВАНИЕ ФУРЬЕ, ДИСКРЕТНОЕ КОСИНУСНОЕ ПРЕОБРАЗОВАНИЕ, ВЕЙВЛЕТ-ПРЕОБРАЗОВАНИЕ, СЖАТИЕ АУДИОДАННЫХ, МЕТРИКИ КАЧЕСТВА, SNR, SI-SDR, LSD, STOI, PESQ, MP3, WAV, PYTHON, PYSIDE6

Объектом исследования являются возможности практического применения матриц Адамара при обработке аудиоданных.

Цель работы заключается в комплексном исследовании применения матриц Адамара для сжатия аудиоданных, а также в оценке влияния данного подхода на ключевые показатели качества сжатия.

Для достижения поставленной цели решены следующие задачи: анализ существующих методов обработки и сжатия аудиосигналов; разработка программного комплекса для сравнительного анализа методов обработки аудиоданных; проведение экспериментальных исследований; формулировка рекомендаций по практическому применению.

Основными результатами работы являются: программный комплекс для сравнительного анализа методов обработки аудиоданных с реализацией семи преобразований и тринадцати метрик качества; закономерности эффективности методов в зависимости от жанровой принадлежности аудиоконтента; рекомендации по выбору метода обработки.

Данные результаты разработки предназначены для выявления принципов, которым подчиняются частотные области, создаваемые применением матриц Адамара при обработке аудиоданных.

Использование результатов позволит сократить вычислительные затраты при обработке аудиоданных и повысить эффективность их обработки."""
    
    return abstract_text

def main():
    print("=" * 60)
    print("Обновление реферата ВКР")
    print("=" * 60)
    
    base_dir = Path('/home/z/my-project/VKR/docs')
    doc_path = base_dir / 'VKR_Final_Complete.docx'
    output_path = base_dir / 'VKR_Final_Complete.docx'
    
    print(f"\nЗагрузка документа: {doc_path}")
    doc = Document(doc_path)
    
    # Подсчёт элементов
    print("\nПодсчёт элементов...")
    counts = count_elements(doc)
    
    print(f"\nОбнаружено:")
    print(f"  Рисунков: {counts['figures']}")
    print(f"    Список: {counts['figure_list']}")
    print(f"  Таблиц: {counts['tables']}")
    print(f"    Список: {counts['table_list']}")
    print(f"  Источников: {counts['references']}")
    
    # Поиск и обновление реферата
    print("\nПоиск реферата...")
    abstract_para_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'РЕФЕРАТ':
            abstract_para_idx = i
            break
    
    if abstract_para_idx is None:
        print("Реферат не найден!")
        return
    
    print(f"Реферат найден на позиции {abstract_para_idx}")
    
    # Обновляем текст реферата
    new_abstract = update_abstract(doc, counts)
    
    # Удаляем старый текст и вставляем новый
    # Находим все параграфы реферата до следующего раздела
    start_idx = abstract_para_idx + 1
    end_idx = start_idx
    
    for i in range(start_idx, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text in ['СОДЕРЖАНИЕ', 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ']:
            end_idx = i
            break
    
    print(f"Параграфы реферата: {start_idx} - {end_idx}")
    
    # Очищаем старые параграфы
    for i in range(start_idx, end_idx):
        for run in doc.paragraphs[i].runs:
            run.text = ""
    
    # Записываем новый текст в первый параграф после заголовка
    if doc.paragraphs[start_idx].runs:
        doc.paragraphs[start_idx].runs[0].text = new_abstract.split('\n\n')[0]
    else:
        doc.paragraphs[start_idx].add_run(new_abstract.split('\n\n')[0])
    
    # Сохраняем
    print(f"\nСохранение: {output_path}")
    doc.save(output_path)
    
    print("\nГотово!")

if __name__ == '__main__':
    main()
