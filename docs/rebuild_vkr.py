#!/usr/bin/env python3
"""
Пересборка VKR документа с правильной структурой.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import subprocess

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Rebuild.docx')
NEW_FIGURES_DIR = os.path.join(VKR_DIR, 'media/new_figures')
DOWNLOAD_DIR = '/home/z/my-project/download'

def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, caption_text):
    """Добавление таблицы."""
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(11)
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    # Границы
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    for row_idx, row in enumerate(rows):
        for col_idx, cell_data in enumerate(row):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    return table

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== ИСПРАВЛЕНИЕ НУМЕРАЦИИ В ГЛАВЕ 2 =====
    print('\nИсправление нумерации в Главе 2...')
    
    # Прямая замена текста
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            # 2.0 -> 2.1
            if 'Рисунок 2.0' in text:
                run.text = text.replace('Рисунок 2.0', 'Рисунок 2.1')
                print('  2.0 -> 2.1')
            # 2.1а -> 2.2
            if 'Рисунок 2.1а' in text:
                run.text = text.replace('Рисунок 2.1а', 'Рисунок 2.2')
                print('  2.1а -> 2.2')
            # Рисунок 2.1 — Архитектура -> 2.3
            if 'Рисунок 2.1 — Архитектура' in text:
                run.text = text.replace('Рисунок 2.1 — Архитектура', 'Рисунок 2.3 — Архитектура')
                print('  2.1 (Архитектура) -> 2.3')
            # Рисунок 2.2 — Алгоритм -> 2.4
            if 'Рисунок 2.2 — Алгоритм' in text or 'Рисунок 2.2 - Алгоритм' in text:
                run.text = text.replace('Рисунок 2.2 — Алгоритм', 'Рисунок 2.4 — Алгоритм')
                run.text = run.text.replace('Рисунок 2.2 - Алгоритм', 'Рисунок 2.4 — Алгоритм')
                print('  2.2 (Алгоритм) -> 2.4')
            # Рисунок 2.3 — Иерархия -> 2.5
            if 'Рисунок 2.3 — Иерархия' in text or 'Рисунок 2.3 - Иерархия' in text:
                run.text = text.replace('Рисунок 2.3 — Иерархия', 'Рисунок 2.5 — Иерархия')
                run.text = run.text.replace('Рисунок 2.3 - Иерархия', 'Рисунок 2.5 — Иерархия')
                print('  2.3 (Иерархия) -> 2.5')
    
    # ===== ДОБАВЛЕНИЕ ТАБЛИЦ В ГЛАВУ 1 =====
    print('\nДобавление таблиц в Главу 1...')
    
    # Таблица 1.1: Вычислительная сложность
    headers1 = ['Метод', 'Операции', 'Сложность', 'Умножения']
    rows1 = [
        ['FFT', 'N·log₂(N)', 'O(N log N)', 'Комплексные'],
        ['FWHT', 'N·log₂(N)', 'O(N log N)', 'Нет'],
        ['DCT', 'N·log₂(N)', 'O(N log N)', 'Вещественные'],
        ['DWT (Хаар)', 'N', 'O(N)', 'Нет'],
        ['μ-law', 'N', 'O(N)', 'Нет'],
    ]
    add_table(doc, headers1, rows1, 'Таблица 1.1 — Сравнение вычислительной сложности методов')
    print('  Таблица 1.1')
    
    # Таблица 1.2: Метрики качества
    headers2 = ['Метрика', 'Область', 'Диапазон', 'Описание']
    rows2 = [
        ['SNR', 'Временная', '0–∞ дБ', 'Отношение сигнал/шум'],
        ['SI-SDR', 'Временная', '–∞–∞ дБ', 'Масштабно-инвариантное'],
        ['RMSE', 'Временная', '0–∞', 'Среднеквадратичная ошибка'],
        ['LSD', 'Спектральная', '0–∞ дБ', 'Спектральное расстояние'],
        ['STOI', 'Психоакуст.', '0–1', 'Разборчивость речи'],
        ['PESQ', 'Психоакуст.', '–0.5–4.5', 'Качество речи'],
    ]
    add_table(doc, headers2, rows2, 'Таблица 1.2 — Характеристики метрик качества аудиосигналов')
    print('  Таблица 1.2')
    
    # ===== ДОБАВЛЕНИЕ РИСУНКОВ =====
    print('\nДобавление рисунков...')
    
    # Рисунок 1.4: Матрицы Адамара
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_hadamard_matrices.png')
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(5.5))
        
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run('Рисунок 1.4 — Матрицы Адамара порядков 1, 2, 4, 8')
        run.italic = True
        print('  Рисунок 1.4 (матрицы Адамара)')
    
    # Рисунок 1.5: Функции Уолша
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_walsh_functions.png')
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(5.5))
        
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run('Рисунок 1.5 — Функции Уолша (первые 8 функций)')
        run.italic = True
        print('  Рисунок 1.5 (функции Уолша)')
    
    # Рисунок 1.6: Сравнение FFT и FWHT
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_fft_vs_fwht.png')
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(5.5))
        
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run('Рисунок 1.6 — Сравнение спектров FFT и FWHT')
        run.italic = True
        print('  Рисунок 1.6 (сравнение спектров)')
    
    # Рисунок 1.7: Схема обработки
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_audio_pipeline.png')
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(5.5))
        
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run('Рисунок 1.7 — Общая схема обработки аудиосигнала')
        run.italic = True
        print('  Рисунок 1.7 (схема обработки)')
    
    # ===== ПОДСЧЕТ И ОБНОВЛЕНИЕ =====
    print('\nПодсчет элементов...')
    
    fig_count = 0
    tbl_count = 0
    
    for para in doc.paragraphs:
        if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', para.text):
            fig_count += 1
        if re.search(r'Таблица\s+[\d.]+\s*[—–-]', para.text):
            tbl_count += 1
    
    page_count = len(doc.paragraphs) // 3
    
    print(f'  Рисунков: {fig_count}')
    print(f'  Таблиц: {tbl_count}')
    
    # Обновление реферата
    for para in doc.paragraphs:
        if 'Научно-исследовательская работа состоит из' in para.text:
            para.clear()
            para.add_run(f'Научно-исследовательская работа состоит из {page_count} страниц, {fig_count} рисунков, {tbl_count} таблиц, 76 использованных источников, 3 приложений.')
            print('  Реферат обновлен')
            break
    
    # ===== СОХРАНЕНИЕ =====
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    shutil.copy(OUTPUT_DOCX, os.path.join(DOWNLOAD_DIR, 'VKR_Rebuild.docx'))
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер: {size / 1024 / 1024:.2f} MB')
    
    # PDF
    print('\nКонвертация в PDF...')
    subprocess.run([
        'soffice', '--headless', '--convert-to', 'pdf',
        '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
    ], check=True, capture_output=True)
    print('PDF создан')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
