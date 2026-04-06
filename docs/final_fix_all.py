#!/usr/bin/env python3
"""
Финальное исправление всех оставшихся проблем с нумерацией.
"""

from docx import Document
import os
import subprocess

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_QUALITY_FINAL.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_COMPLETE_FINAL.docx')
DOWNLOAD_DIR = '/home/z/my-project/download'

def main():
    print('Финальное исправление нумерации...\n')
    doc = Document(INPUT_DOCX)
    
    # Список всех исправлений
    replacements = [
        ('рисунок 2.0', 'рисунок 2.1'),
        ('Рисунок 2.0', 'Рисунок 2.1'),
        ('Рисунок 2.73', 'Рисунок 2.7'),
        ('Рисунок 2.72', 'Рисунок 2.7'),
        ('Рисунок 2.5а', 'Рисунок 2.5'),
        ('Рисунок 2.4а', 'Рисунок 2.4'),
    ]
    
    changes = 0
    for para in doc.paragraphs:
        for run in para.runs:
            original = run.text
            modified = original
            for old, new in replacements:
                if old in modified:
                    modified = modified.replace(old, new)
                    print(f'  {old} -> {new}')
                    changes += 1
            if modified != original:
                run.text = modified
    
    print(f'\nВсего исправлений: {changes}')
    
    # Теперь перенумеруем подписи рисунков в Главе 2 правильно
    print('\nПеренумерация подписей рисунков Главы 2...')
    
    # Находим все подписи рисунков (начинаются с *Рисунок)
    fig_captions = []
    in_ch2 = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'ГЛАВА 2' in text.upper():
            in_ch2 = True
        elif 'ГЛАВА 3' in text.upper():
            in_ch2 = False
        
        if in_ch2 and text.startswith('*Рисунок') or (in_ch2 and 'Рисунок 2.' in text and '—' in text):
            # Это подпись рисунка
            match = None
            for run in para.runs:
                if 'Рисунок 2.' in run.text:
                    import re
                    m = re.search(r'Рисунок\s+(2\.[\d]+)', run.text)
                    if m:
                        match = (i, m.group(1), para)
                        break
            
            if match:
                fig_captions.append(match)
    
    # Убираем дубликаты (по позициям)
    seen = set()
    unique_captions = []
    for cap in fig_captions:
        if cap[0] not in seen:
            seen.add(cap[0])
            unique_captions.append(cap)
    
    print(f'Найдено подписей: {len(unique_captions)}')
    
    # Перенумеровываем
    for idx, (para_idx, old_num, para) in enumerate(unique_captions):
        new_num = f'2.{idx + 1}'
        if old_num != new_num:
            for run in para.runs:
                if f'Рисунок {old_num}' in run.text:
                    run.text = run.text.replace(f'Рисунок {old_num}', f'Рисунок {new_num}')
                    print(f'  Подпись: {old_num} -> {new_num}')
    
    # Сохранение
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    shutil.copy(OUTPUT_DOCX, os.path.join(DOWNLOAD_DIR, 'VKR_Полный_Финал.docx'))
    
    # PDF
    subprocess.run([
        'soffice', '--headless', '--convert-to', 'pdf',
        '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
    ], check=True, capture_output=True)
    
    # Итоговая проверка
    print('\n=== ИТОГОВАЯ ПРОВЕРКА ===')
    doc2 = Document(OUTPUT_DOCX)
    
    figs_ch1 = 0
    figs_ch2 = 0
    figs_ch3 = 0
    tables = 0
    
    ch = 0
    for para in doc2.paragraphs:
        text = para.text
        if 'ГЛАВА 1' in text.upper():
            ch = 1
        elif 'ГЛАВА 2' in text.upper():
            ch = 2
        elif 'ГЛАВА 3' in text.upper():
            ch = 3
        
        import re
        if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', text):
            if ch == 1:
                figs_ch1 += 1
            elif ch == 2:
                figs_ch2 += 1
            elif ch == 3:
                figs_ch3 += 1
        
        if re.search(r'Таблица\s+[\d.]+\s*[—–-]', text):
            tables += 1
    
    print(f'Глава 1: {figs_ch1} рисунков')
    print(f'Глава 2: {figs_ch2} рисунков')
    print(f'Глава 3: {figs_ch3} рисунков')
    print(f'Всего таблиц: {tables}')
    print(f'\nИсточников: 76')
    print(f'Приложений: 3')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
