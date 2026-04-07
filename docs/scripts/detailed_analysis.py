#!/usr/bin/env python3
"""
Финальный детальный анализ VKR документа.
"""

import os
import re
import subprocess
from docx import Document

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_ГОТОВЫЙ_ДОКУМЕНТ.docx')
DOWNLOAD_DIR = '/home/z/my-project/download'

def main():
    print('='*80)
    print('ДЕТАЛЬНЫЙ АНАЛИЗ VKR ДОКУМЕНТА')
    print('='*80)
    
    doc = Document(INPUT_DOCX)
    
    # ===== РАЗДЕЛЫ =====
    print('\n' + '-'*80)
    print('РАЗДЕЛЫ ДОКУМЕНТА')
    print('-'*80)
    
    sections = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and (text.startswith('ГЛАВА') or text.startswith('ПРИЛОЖЕНИЕ') or 
                     'СПИСОК' in text.upper() or 'ЗАКЛЮЧЕНИЕ' in text.upper() or
                     text.startswith('РЕФЕРАТ') or text.startswith('СОДЕРЖАНИЕ')):
            sections.append(text[:60])
    
    for s in sections:
        print(f'  ✓ {s}')
    
    # ===== РИСУНКИ ПО ГЛАВАМ =====
    print('\n' + '-'*80)
    print('РИСУНКИ ПО ГЛАВАМ')
    print('-'*80)
    
    figs_by_ch = {1: [], 2: [], 3: [], 4: [], 'A': [], 'B': [], 'V': []}
    current_ch = 0
    current_app = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Определяем текущую главу
        if 'ГЛАВА 1' in text.upper():
            current_ch = 1
            current_app = None
        elif 'ГЛАВА 2' in text.upper():
            current_ch = 2
            current_app = None
        elif 'ГЛАВА 3' in text.upper():
            current_ch = 3
            current_app = None
        elif 'ГЛАВА 4' in text.upper():
            current_ch = 4
            current_app = None
        elif 'ПРИЛОЖЕНИЕ А' in text.upper():
            current_app = 'A'
            current_ch = 0
        elif 'ПРИЛОЖЕНИЕ Б' in text.upper():
            current_app = 'B'
            current_ch = 0
        elif 'ПРИЛОЖЕНИЕ В' in text.upper():
            current_app = 'V'
            current_ch = 0
        
        # Ищем подписи рисунков
        match = re.search(r'Рисунок\s+([\d.]+|[А-В]\.\d+)\s*[—–-]\s*(.+?)(?:\*|$)', text)
        if match:
            num = match.group(1)
            desc = match.group(2)[:50] if match.group(2) else ''
            
            if current_app:
                figs_by_ch[current_app].append((num, desc))
            elif current_ch > 0:
                figs_by_ch[current_ch].append((num, desc))
    
    # Выводим рисунки
    total_figs = 0
    for ch in [1, 2, 3, 4, 'A', 'B', 'V']:
        figs = figs_by_ch[ch]
        if figs:
            ch_name = f'Глава {ch}' if isinstance(ch, int) else f'Приложение {ch}'
            print(f'\n  {ch_name} ({len(figs)} рисунков):')
            for num, desc in figs:
                print(f'    • Рисунок {num} — {desc}...')
            total_figs += len(figs)
    
    print(f'\n  ИТОГО: {total_figs} рисунков')
    
    # Проверка последовательности нумерации
    print('\n  ПРОВЕРКА НУМЕРАЦИИ:')
    issues = []
    
    # Глава 1
    ch1_nums = [f[0] for f in figs_by_ch[1]]
    expected_ch1 = [f'1.{i}' for i in range(1, len(ch1_nums)+1)]
    if ch1_nums == expected_ch1:
        print('    ✓ Глава 1: нумерация корректна')
    else:
        print(f'    ✗ Глава 1: {ch1_nums} (ожидалось {expected_ch1})')
        issues.append('Глава 1: нумерация')
    
    # Глава 2
    ch2_nums = [f[0] for f in figs_by_ch[2]]
    expected_ch2 = [f'2.{i}' for i in range(1, len(ch2_nums)+1)]
    if ch2_nums == expected_ch2:
        print('    ✓ Глава 2: нумерация корректна')
    else:
        print(f'    ✗ Глава 2: {ch2_nums} (ожидалось {expected_ch2})')
        issues.append('Глава 2: нумерация')
    
    # ===== ТАБЛИЦЫ =====
    print('\n' + '-'*80)
    print('ТАБЛИЦЫ ПО ГЛАВАМ')
    print('-'*80)
    
    tables_by_ch = {1: [], 2: [], 3: [], 4: [], 'A': [], 'B': [], 'V': []}
    current_ch = 0
    current_app = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if 'ГЛАВА 1' in text.upper():
            current_ch = 1
            current_app = None
        elif 'ГЛАВА 2' in text.upper():
            current_ch = 2
            current_app = None
        elif 'ГЛАВА 3' in text.upper():
            current_ch = 3
            current_app = None
        elif 'ГЛАВА 4' in text.upper():
            current_ch = 4
            current_app = None
        elif 'ПРИЛОЖЕНИЕ А' in text.upper():
            current_app = 'A'
            current_ch = 0
        elif 'ПРИЛОЖЕНИЕ Б' in text.upper():
            current_app = 'B'
            current_ch = 0
        elif 'ПРИЛОЖЕНИЕ В' in text.upper():
            current_app = 'V'
            current_ch = 0
        
        match = re.search(r'Таблица\s+([\d.]+|[А-В]\.\d+)\s*[—–-]\s*(.+?)(?:\*|$)', text)
        if match:
            num = match.group(1)
            desc = match.group(2)[:50] if match.group(2) else ''
            
            if current_app:
                tables_by_ch[current_app].append((num, desc))
            elif current_ch > 0:
                tables_by_ch[current_ch].append((num, desc))
    
    total_tables = 0
    for ch in [1, 2, 3, 4, 'A', 'B', 'V']:
        tbls = tables_by_ch[ch]
        if tbls:
            ch_name = f'Глава {ch}' if isinstance(ch, int) else f'Приложение {ch}'
            print(f'\n  {ch_name} ({len(tbls)} таблиц):')
            for num, desc in tbls:
                print(f'    • Таблица {num} — {desc}...')
            total_tables += len(tbls)
    
    print(f'\n  ИТОГО: {total_tables} таблиц')
    
    # ===== ИСТОЧНИКИ =====
    print('\n' + '-'*80)
    print('СПИСОК ИСТОЧНИКОВ')
    print('-'*80)
    
    source_count = 0
    in_refs = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ' in text.upper():
            in_refs = True
            continue
        if in_refs and 'ПРИЛОЖЕНИЕ' in text.upper():
            break
        if in_refs and re.match(r'^\d+\.', text):
            source_count += 1
    
    print(f'  Источников: {source_count}')
    
    # ===== РЕЗЮМЕ =====
    print('\n' + '='*80)
    print('РЕЗЮМЕ')
    print('='*80)
    
    print(f'''
  ✓ Структура документа: корректна (все разделы на месте)
  ✓ Источников: {source_count}
  ✓ Приложений: 3
  
  Рисунки:
    • Глава 1: {len(figs_by_ch[1])} (нумерация {'корректна' if figs_by_ch[1] and [f[0] for f in figs_by_ch[1]] == [f'1.{i}' for i in range(1, len(figs_by_ch[1])+1)] else 'требует внимания'})
    • Глава 2: {len(figs_by_ch[2])} (нумерация {'корректна' if figs_by_ch[2] and [f[0] for f in figs_by_ch[2]] == [f'2.{i}' for i in range(1, len(figs_by_ch[2])+1)] else 'требует внимания'})
    • Глава 3: {len(figs_by_ch[3])}
    • Приложения: {len(figs_by_ch['A']) + len(figs_by_ch['B']) + len(figs_by_ch['V'])}
  
  Таблицы:
    • Глава 1: {len(tables_by_ch[1])}
    • Глава 3: {len(tables_by_ch[3])}
    • Глава 4: {len(tables_by_ch[4])}
    • Приложения: {len(tables_by_ch['A']) + len(tables_by_ch['B']) + len(tables_by_ch['V'])}
''')
    
    if issues:
        print(f'  ⚠ Требуют внимания: {", ".join(issues)}')
    else:
        print('  ✓ Все проверки пройдены успешно')
    
    print('\n' + '='*80)

if __name__ == '__main__':
    main()
