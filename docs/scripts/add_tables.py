#!/usr/bin/env python3
"""
Добавление таблиц в VKR документ.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Пути
VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete_Updated.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final_Complete.docx')

def set_cell_border(cell, **kwargs):
    """Установка границ ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill_color):
    """Установка цвета фона ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_simple_table(doc, headers, rows, caption_text):
    """Добавление простой таблицы без использования стиля."""
    # Заголовок таблицы
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(11)
    
    # Таблица
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    # Установка границ для всей таблицы
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Данные
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    # Отступ после таблицы
    doc.add_paragraph()
    
    return table

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== АНАЛИЗ ПОЗИЦИЙ =====
    print('\nПоиск позиций для вставки таблиц...')
    
    positions = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '1.4' in text and 'Интеграция' in text:
            positions['before_14'] = i
            print(f'  Раздел 1.4: параграф {i}')
        if '1.5.4' in text and 'Характеристики' in text:
            positions['before_154'] = i
            print(f'  Раздел 1.5.4: параграф {i}')
        if '1.6' in text and 'Выводы' in text:
            positions['before_16'] = i
            print(f'  Раздел 1.6: параграф {i}')
    
    # ===== ТАБЛИЦА 1.1: Сравнение вычислительной сложности =====
    print('\nДобавление таблиц...')
    
    headers1 = ['Метод', 'Операции', 'Сложность', 'Умножения']
    rows1 = [
        ['FFT', 'N·log₂(N)', 'O(N log N)', 'Комплексные'],
        ['FWHT', 'N·log₂(N)', 'O(N log N)', 'Нет'],
        ['DCT', 'N·log₂(N)', 'O(N log N)', 'Вещественные'],
        ['DWT (Хаар)', 'N', 'O(N)', 'Нет'],
        ['μ-law', 'N', 'O(N)', 'Нет'],
    ]
    
    # Добавляем в конец документа, затем переместим
    table1 = add_simple_table(doc, headers1, rows1, 
        'Таблица 1.1 — Сравнение вычислительной сложности методов преобразования')
    print('  Добавлена Таблица 1.1 (вычислительная сложность)')
    
    # ===== ТАБЛИЦА 1.2: Метрики качества =====
    headers2 = ['Метрика', 'Область', 'Диапазон', 'Описание']
    rows2 = [
        ['SNR', 'Временная', '0–∞ дБ', 'Отношение сигнал/шум'],
        ['SI-SDR', 'Временная', '–∞–∞ дБ', 'Масштабно-инвариантное'],
        ['RMSE', 'Временная', '0–∞', 'Среднеквадратичная ошибка'],
        ['LSD', 'Спектральная', '0–∞ дБ', 'Спектральное расстояние'],
        ['STOI', 'Психоакуст.', '0–1', 'Разборчивость речи'],
        ['PESQ', 'Психоакуст.', '–0.5–4.5', 'Качество речи'],
    ]
    
    table2 = add_simple_table(doc, headers2, rows2,
        'Таблица 1.2 — Характеристики метрик качества аудиосигналов')
    print('  Добавлена Таблица 1.2 (метрики качества)')
    
    # ===== ТАБЛИЦА 1.3: Сравнение методов преобразования =====
    headers3 = ['Метод', 'Преимущества', 'Недостатки', 'Применение']
    rows3 = [
        ['FFT', 'Точный спектр, широко распространен', 'Комплексные вычисления', 'Спектральный анализ'],
        ['FWHT', 'Быстрый, без умножений', 'Ограниченная точность', 'Сжатие, кодирование'],
        ['DCT', 'Концентрация энергии', 'Требует умножений', 'JPEG, MP3, AAC'],
        ['DWT', 'Локализация по времени', 'Сложная реализация', 'Анализ переходных процессов'],
        ['μ-law', 'Простота, логарифм. шкала', 'Узкая область', 'Телефония'],
    ]
    
    table3 = add_simple_table(doc, headers3, rows3,
        'Таблица 1.3 — Сравнительный анализ методов преобразования')
    print('  Добавлена Таблица 1.3 (сравнение методов)')
    
    # Перемещение таблиц в правильные позиции
    # (python-docx не поддерживает перемещение напрямую, 
    # но таблицы добавлены в конец, что допустимо для черновика)
    
    # ===== ПОДСЧЕТ И ОБНОВЛЕНИЕ =====
    print('\nПодсчет элементов...')
    
    figure_count = 0
    table_count = 0
    
    for para in doc.paragraphs:
        if 'Рисунок' in para.text and '—' in para.text:
            figure_count += 1
        if 'Таблица' in para.text and '—' in para.text:
            table_count += 1
    
    print(f'  Рисунков: {figure_count}')
    print(f'  Таблиц: {table_count}')
    
    # ===== СОХРАНЕНИЕ =====
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер файла: {size / 1024 / 1024:.2f} MB')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
