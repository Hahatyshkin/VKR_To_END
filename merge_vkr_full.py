#!/usr/bin/env python3
"""
Полное объединение ВКР с расширенными главами и встроенными изображениями.
Создаёт новый документ с нуля, объединяя все части.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from pathlib import Path

def copy_paragraph(source_para, target_doc):
    """Копирование параграфа из одного документа в другой"""
    new_para = target_doc.add_paragraph()
    new_para.style = source_para.style
    new_para.alignment = source_para.alignment
    
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        if run.bold:
            new_run.bold = True
        if run.italic:
            new_run.italic = True
        if run.underline:
            new_run.underline = True
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
    
    return new_para

def copy_table(source_table, target_doc):
    """Копирование таблицы из одного документа в другой"""
    # Создаём новую таблицу с тем же количеством строк и столбцов
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    new_table = target_doc.add_table(rows=rows, cols=cols)
    new_table.style = source_table.style
    
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            for para in cell.paragraphs:
                new_para = new_cell.add_paragraph()
                new_para.style = para.style
                new_para.alignment = para.alignment
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    if run.bold:
                        new_run.bold = True
                    if run.italic:
                        new_run.italic = True
                    if run.font.size:
                        new_run.font.size = run.font.size
    
    return new_table

def find_element_index(doc, search_text, element_type='paragraph'):
    """Поиск индекса элемента по тексту"""
    if element_type == 'paragraph':
        for i, para in enumerate(doc.paragraphs):
            if search_text.lower() in para.text.lower():
                return i
    return None

def get_section_boundaries(doc):
    """Определение границ разделов документа"""
    boundaries = {
        'ch3_start': None,
        'ch3_end': None,
        'ch4_start': None,
        'ch4_end': None,
        'conclusion_start': None,
        'references_start': None,
        'appendix_a_start': None,
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'ГЛАВА 3' in text or '[ГЛАВА 3' in text:
            boundaries['ch3_start'] = i
        elif 'ГЛАВА 4' in text or '[ГЛАВА 4' in text:
            boundaries['ch3_end'] = i
            boundaries['ch4_start'] = i
        elif text.startswith('ЗАКЛЮЧЕНИЕ') or text.startswith('[ЗАКЛЮЧЕНИЕ'):
            boundaries['ch4_end'] = i
            boundaries['conclusion_start'] = i
        elif 'СПИСОК ИСПОЛЬЗОВАННЫХ' in text:
            boundaries['references_start'] = i
        elif 'ПРИЛОЖЕНИЕ А' in text:
            boundaries['appendix_a_start'] = i
    
    return boundaries

def add_graphs_to_chapter(doc, graphs_dir, chapter_num):
    """Добавление графиков в главу"""
    graph_files = {
        3: [
            ('snr_comparison.png', 'Рисунок 3.1 --- Сравнение методов по SNR'),
            ('sisdr_comparison.png', 'Рисунок 3.2 --- Сравнение методов по SI-SDR'),
            ('lsd_comparison.png', 'Рисунок 3.3 --- Сравнение методов по LSD'),
            ('time_comparison.png', 'Рисунок 3.4 --- Сравнение методов по времени обработки'),
            ('size_comparison.png', 'Рисунок 3.5 --- Сравнение методов по размеру'),
            ('radar_comparison.png', 'Рисунок 3.6 --- Радарная диаграмма методов'),
            ('snr_vs_time.png', 'Рисунок 3.7 --- Соотношение качество/скорость'),
            ('snr_vs_size.png', 'Рисунок 3.8 --- Соотношение качество/размер'),
            ('correlation_heatmap.png', 'Рисунок 3.9 --- Корреляционная матрица'),
            ('snr_boxplot.png', 'Рисунок 3.10 --- Распределение SNR'),
        ],
        4: [
            ('snr_comparison_all.png', 'Рисунок 4.1 --- Сравнение методов по SNR'),
            ('time_comparison_all.png', 'Рисунок 4.2 --- Сравнение по времени'),
            ('radar_comparison.png', 'Рисунок 4.3 --- Радарная диаграмма'),
            ('snr_vs_time.png', 'Рисунок 4.4 --- Компромисс качество/скорость'),
            ('lsd_comparison.png', 'Рисунок 4.5 --- Сравнение по LSD'),
        ]
    }
    
    added = 0
    for filename, caption in graph_files.get(chapter_num, []):
        filepath = graphs_dir / filename
        if filepath.exists():
            # Добавляем параграф с изображением
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            try:
                run.add_picture(str(filepath), width=Inches(5.5))
                # Добавляем подпись
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption)
                caption_run.italic = True
                caption_run.font.size = Pt(10)
                added += 1
                print(f"    Добавлено: {filename}")
            except Exception as e:
                print(f"    Ошибка добавления {filename}: {e}")
    
    return added

def main():
    print("=" * 60)
    print("Объединение ВКР с расширенными главами")
    print("=" * 60)
    
    base_dir = Path('/home/z/my-project/VKR')
    docs_dir = base_dir / 'docs'
    final_dir = docs_dir / 'final'
    output_path = docs_dir / 'VKR_Complete.docx'
    
    graphs_ch3 = final_dir / 'graphs_chapter3'
    graphs_ch4 = final_dir / 'graphs_chapter4'
    
    # Загружаем документы
    print("\n1. Загрузка документов...")
    main_doc = Document(docs_dir / 'VKR.docx')
    ch3_doc = Document(final_dir / 'VKR_Chapter3_Extended.docx')
    ch4_doc = Document(final_dir / 'VKR_Chapter4_Extended.docx')
    
    # Определяем границы разделов
    print("2. Определение структуры основного документа...")
    boundaries = get_section_boundaries(main_doc)
    print(f"   Глава 3: {boundaries['ch3_start']} - {boundaries['ch3_end']}")
    print(f"   Глава 4: {boundaries['ch4_start']} - {boundaries['ch4_end']}")
    print(f"   Заключение: {boundaries['conclusion_start']}")
    print(f"   Список источников: {boundaries['references_start']}")
    print(f"   Приложение А: {boundaries['appendix_a_start']}")
    
    # Создаём новый документ
    print("\n3. Создание объединённого документа...")
    result_doc = Document()
    
    # Копируем стили из основного документа (простое копирование)
    # python-docx не поддерживает прямое копирование стилей, 
    # поэтому копируем параграфы с сохранением форматирования
    
    # Часть 1: Всё до главы 3
    print("   Копирование: Титульный лист, Реферат, Содержание, Введение, Главы 1-2...")
    
    # Получаем все элементы (параграфы и таблицы) в порядке их появления
    body = main_doc.element.body
    
    for i, para in enumerate(main_doc.paragraphs):
        if boundaries['ch3_start'] and i >= boundaries['ch3_start']:
            break
        copy_paragraph(para, result_doc)
    
    # Часть 2: Расширенная глава 3 с графиками
    print("   Добавление расширенной главы 3 с графиками...")
    for para in ch3_doc.paragraphs:
        copy_paragraph(para, result_doc)
    
    # Добавляем графики главы 3
    print("   Вставка графиков главы 3:")
    add_graphs_to_chapter(result_doc, graphs_ch3, 3)
    
    # Часть 3: Расширенная глава 4 с графиками
    print("   Добавление расширенной главы 4 с графиками...")
    for para in ch4_doc.paragraphs:
        copy_paragraph(para, result_doc)
    
    # Добавляем графики главы 4
    print("   Вставка графиков главы 4:")
    add_graphs_to_chapter(result_doc, graphs_ch4, 4)
    
    # Часть 4: Заключение, Список источников, Приложения
    print("   Копирование: Заключение, Список источников, Приложения...")
    
    if boundaries['conclusion_start']:
        for i, para in enumerate(main_doc.paragraphs):
            if i >= boundaries['conclusion_start']:
                copy_paragraph(para, result_doc)
    
    # Сохраняем результат
    print(f"\n4. Сохранение: {output_path}")
    result_doc.save(output_path)
    
    # Проверяем размер
    size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"\nГотово! Размер файла: {size_mb:.2f} MB")
    
    return output_path

if __name__ == '__main__':
    main()
