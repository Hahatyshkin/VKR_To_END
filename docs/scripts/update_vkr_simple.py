#!/usr/bin/env python3
"""
Обновление VKR с использованием python-docx для более надежного редактирования.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from copy import deepcopy

# Пути
VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final.docx')
NEW_FIGURES_DIR = os.path.join(VKR_DIR, 'media/new_figures')

def set_cell_shading(cell, fill_color):
    """Установка цвета фона ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_figure(doc, image_path, caption_text, width_inches=5.5):
    """Добавление рисунка с подписью."""
    # Параграф для изображения
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    
    # Подпись
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(11)
    
    return p

def add_table_with_caption(doc, headers, rows, caption_text):
    """Добавление таблицы с заголовком."""
    # Заголовок таблицы
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.bold = True
    
    # Таблица
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9D9D9')
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    
    # Данные
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    # Пустой параграф после таблицы
    doc.add_paragraph()
    
    return table

def find_paragraph_index(doc, search_text):
    """Найти индекс параграфа по тексту."""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            return i
    return -1

def find_heading_index(doc, heading_text):
    """Найти индекс заголовка по тексту."""
    for i, para in enumerate(doc.paragraphs):
        if heading_text.lower() in para.text.lower():
            # Проверяем что это заголовок
            if para.style and 'heading' in para.style.name.lower():
                return i
    return -1

def update_figure_numbers_chapter2(doc):
    """Исправление нумерации рисунков в Главе 2."""
    changes = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Исправляем "Рисунок 2.0" -> "Рисунок 2.1"
        if 'Рисунок 2.0' in text:
            for run in para.runs:
                if 'Рисунок 2.0' in run.text:
                    run.text = run.text.replace('Рисунок 2.0', 'Рисунок 2.1')
                    changes.append(f'Строка {i}: 2.0 -> 2.1')
        
        # Исправляем "Рисунок 2.1а" -> "Рисунок 2.2"
        if 'Рисунок 2.1а' in text:
            for run in para.runs:
                if 'Рисунок 2.1а' in run.text:
                    run.text = run.text.replace('Рисунок 2.1а', 'Рисунок 2.2')
                    changes.append(f'Строка {i}: 2.1а -> 2.2')
        
        # Исправляем "Рисунок 2.1 —" -> "Рисунок 2.3 —"
        if 'Рисунок 2.1 —' in text or 'Рисунок 2.1 -' in text:
            for run in para.runs:
                if 'Рисунок 2.1 —' in run.text or 'Рисунок 2.1 -' in run.text:
                    run.text = run.text.replace('Рисунок 2.1 —', 'Рисунок 2.3 —')
                    run.text = run.text.replace('Рисунок 2.1 -', 'Рисунок 2.3 —')
                    changes.append(f'Строка {i}: 2.1 -> 2.3')
        
        # Исправляем "Рисунок 2.2 —" -> "Рисунок 2.4 —"
        if 'Рисунок 2.2 —' in text or 'Рисунок 2.2 -' in text:
            for run in para.runs:
                if 'Рисунок 2.2 —' in run.text or 'Рисунок 2.2 -' in run.text:
                    run.text = run.text.replace('Рисунок 2.2 —', 'Рисунок 2.4 —')
                    run.text = run.text.replace('Рисунок 2.2 -', 'Рисунок 2.4 —')
                    changes.append(f'Строка {i}: 2.2 -> 2.4')
        
        # Исправляем "Рисунок 2.3 —" -> "Рисунок 2.5 —"
        if 'Рисунок 2.3 —' in text or 'Рисунок 2.3 -' in text:
            for run in para.runs:
                if 'Рисунок 2.3 —' in run.text or 'Рисунок 2.3 -' in run.text:
                    run.text = run.text.replace('Рисунок 2.3 —', 'Рисунок 2.5 —')
                    run.text = run.text.replace('Рисунок 2.3 -', 'Рисунок 2.5 —')
                    changes.append(f'Строка {i}: 2.3 -> 2.5')
    
    return changes

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== ИСПРАВЛЕНИЕ НУМЕРАЦИИ В ГЛАВЕ 2 =====
    print('\nИсправление нумерации рисунков в Главе 2...')
    changes = update_figure_numbers_chapter2(doc)
    for c in changes:
        print(f'  {c}')
    
    # ===== ПОДСЧЕТ ТЕКУЩИХ РИСУНКОВ И ТАБЛИЦ =====
    figure_count = 0
    table_count = 0
    
    for para in doc.paragraphs:
        if 'Рисунок' in para.text:
            figure_count += 1
        if 'Таблица' in para.text:
            table_count += 1
    
    print(f'\nТекущее состояние:')
    print(f'  Рисунков: {figure_count}')
    print(f'  Таблиц: {table_count}')
    
    # ===== ДОБАВЛЕНИЕ ТАБЛИЦ В ГЛАВУ 1 =====
    print('\nДобавление таблиц в Главу 1...')
    
    # Находим конец раздела 1.3 (перед 1.4)
    idx_14 = -1
    idx_15 = -1
    idx_16 = -1
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '1.4' in text and 'Интеграция' in text:
            idx_14 = i
        if '1.5' in text and 'Метрики' in text:
            idx_15 = i
        if '1.6' in text and 'Выводы' in text:
            idx_16 = i
    
    print(f'  Индекс раздела 1.4: {idx_14}')
    print(f'  Индекс раздела 1.5: {idx_15}')
    print(f'  Индекс раздела 1.6: {idx_16}')
    
    # Сохранение документа
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    # Проверка
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер файла: {size / 1024 / 1024:.2f} MB')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
