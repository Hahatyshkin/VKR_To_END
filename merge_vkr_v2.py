#!/usr/bin/env python3
"""
Правильное объединение ВКР с расширенными главами и встроенными изображениями.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

def copy_run(source_run, target_para):
    """Копирование run с форматированием"""
    new_run = target_para.add_run(source_run.text)
    if source_run.bold:
        new_run.bold = True
    if source_run.italic:
        new_run.italic = True
    if source_run.underline:
        new_run.underline = source_run.underline
    if source_run.font.size:
        new_run.font.size = source_run.font.size
    if source_run.font.name:
        new_run.font.name = source_run.font.name
    try:
        if source_run.font.color.rgb:
            new_run.font.color.rgb = source_run.font.color.rgb
    except:
        pass
    return new_run

def copy_para(source_para, target_doc):
    """Копирование параграфа"""
    new_para = target_doc.add_paragraph()
    
    # Копируем выравнивание
    new_para.alignment = source_para.alignment
    
    # Копируем все runs
    for run in source_para.runs:
        copy_run(run, new_para)
    
    # Пробуем скопировать стиль
    try:
        new_para.style = source_para.style
    except:
        pass
    
    # Копируем отступы
    if source_para.paragraph_format.first_line_indent:
        new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
    if source_para.paragraph_format.left_indent:
        new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    if source_para.paragraph_format.space_before:
        new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    if source_para.paragraph_format.space_after:
        new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    
    return new_para

def copy_table_simple(source_table, target_doc):
    """Простое копирование таблицы"""
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    new_table = target_doc.add_table(rows=rows, cols=cols)
    
    try:
        new_table.style = 'Table Grid'
    except:
        pass
    
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            # Копируем текст ячейки
            for para in cell.paragraphs:
                if para.text.strip():
                    new_para = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
                    for run in para.runs:
                        copy_run(run, new_para)
    
    return new_table

def add_graph(doc, image_path, caption):
    """Добавление графика с подписью"""
    if not image_path.exists():
        print(f"    Файл не найден: {image_path}")
        return False
    
    # Параграф с изображением
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    
    try:
        run.add_picture(str(image_path), width=Inches(5.5))
        
        # Подпись
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        
        return True
    except Exception as e:
        print(f"    Ошибка: {e}")
        return False

def main():
    print("=" * 60)
    print("Создание единого документа ВКР")
    print("=" * 60)
    
    base_dir = Path('/home/z/my-project/VKR')
    docs_dir = base_dir / 'docs'
    final_dir = docs_dir / 'final'
    output_path = docs_dir / 'VKR_Final_Complete.docx'
    
    graphs_ch3 = final_dir / 'graphs_chapter3'
    graphs_ch4 = final_dir / 'graphs_chapter4'
    
    # Загружаем все документы
    print("\nЗагрузка документов...")
    main_doc = Document(docs_dir / 'VKR.docx')
    ch3_doc = Document(final_dir / 'VKR_Chapter3_Extended.docx')
    ch4_doc = Document(final_dir / 'VKR_Chapter4_Extended.docx')
    
    # Определяем границы разделов в основном документе
    print("Определение структуры...")
    sections = {}
    for i, para in enumerate(main_doc.paragraphs):
        text = para.text.strip()
        if 'ГЛАВА 1' in text and sections.get('ch1') is None:
            sections['ch1'] = i
        elif 'ГЛАВА 2' in text and sections.get('ch2') is None:
            sections['ch2'] = i
        elif 'ГЛАВА 3' in text and sections.get('ch3') is None:
            sections['ch3'] = i
        elif 'ГЛАВА 4' in text and sections.get('ch4') is None:
            sections['ch4'] = i
        elif text.startswith('ЗАКЛЮЧЕНИЕ'):
            sections['conclusion'] = i
        elif 'СПИСОК ИСПОЛЬЗОВАННЫХ' in text:
            sections['references'] = i
        elif 'ПРИЛОЖЕНИЕ А' in text:
            sections['app_a'] = i
    
    print(f"  Глава 1: {sections.get('ch1')}")
    print(f"  Глава 2: {sections.get('ch2')}")
    print(f"  Глава 3: {sections.get('ch3')}")
    print(f"  Глава 4: {sections.get('ch4')}")
    print(f"  Заключение: {sections.get('conclusion')}")
    print(f"  Список источников: {sections.get('references')}")
    print(f"  Приложение А: {sections.get('app_a')}")
    
    # Создаём новый документ
    print("\nСоздание документа...")
    result_doc = Document()
    
    # Графики для вставки
    ch3_graphs = [
        ('snr_comparison.png', 'Рисунок 3.1 --- Сравнение методов по SNR (среднее ± СКО)'),
        ('sisdr_comparison.png', 'Рисунок 3.2 --- Сравнение методов по SI-SDR'),
        ('lsd_comparison.png', 'Рисунок 3.3 --- Сравнение методов по LSD (ниже = лучше)'),
        ('time_comparison.png', 'Рисунок 3.4 --- Сравнение методов по времени обработки'),
        ('size_comparison.png', 'Рисунок 3.5 --- Сравнение методов по размеру выходного файла'),
        ('radar_comparison.png', 'Рисунок 3.6 --- Радарная диаграмма методов'),
        ('snr_vs_time.png', 'Рисунок 3.7 --- Соотношение качество/скорость'),
        ('snr_vs_size.png', 'Рисунок 3.8 --- Соотношение качество/размер'),
        ('correlation_heatmap.png', 'Рисунок 3.9 --- Корреляционная матрица метрик'),
        ('snr_boxplot.png', 'Рисунок 3.10 --- Распределение SNR по методам'),
    ]
    
    ch4_graphs = [
        ('snr_comparison_all.png', 'Рисунок 4.1 --- Сравнение методов по SNR'),
        ('time_comparison_all.png', 'Рисунок 4.2 --- Сравнение по времени обработки'),
        ('radar_comparison.png', 'Рисунок 4.3 --- Радарная диаграмма методов'),
        ('snr_vs_time.png', 'Рисунок 4.4 --- Компромисс качество/скорость'),
        ('lsd_comparison.png', 'Рисунок 4.5 --- Сравнение по LSD'),
    ]
    
    # ЧАСТЬ 1: До главы 3 (титул, реферат, содержание, введение, главы 1-2)
    print("  Копирование: Титульный лист, Реферат, Содержание, Введение, Главы 1-2...")
    end_part1 = sections.get('ch3', 300)
    for i, para in enumerate(main_doc.paragraphs[:end_part1]):
        copy_para(para, result_doc)
    
    # ЧАСТЬ 2: Расширенная глава 3 с графиками
    print("  Добавление: Глава 3 (расширенная)...")
    graph_idx = 0
    for para in ch3_doc.paragraphs:
        text = para.text.strip()
        copy_para(para, result_doc)
        
        # Вставляем график после соответствующей таблицы/ссылки
        if graph_idx < len(ch3_graphs):
            for j, (fname, caption) in enumerate(ch3_graphs):
                if f'Рисунок 3.{j+1}' in text or f'рисунок 3.{j+1}' in text.lower():
                    # Вставляем график
                    add_graph(result_doc, graphs_ch3 / fname, caption)
                    graph_idx = max(graph_idx, j + 1)
    
    # Добавляем оставшиеся графики главы 3
    print("    Добавление оставшихся графиков главы 3...")
    for fname, caption in ch3_graphs[graph_idx:]:
        add_graph(result_doc, graphs_ch3 / fname, caption)
    
    # ЧАСТЬ 3: Расширенная глава 4 с графиками
    print("  Добавление: Глава 4 (расширенная)...")
    graph_idx = 0
    for para in ch4_doc.paragraphs:
        copy_para(para, result_doc)
        
        # Вставляем график после соответствующей ссылки
        if graph_idx < len(ch4_graphs):
            for j, (fname, caption) in enumerate(ch4_graphs):
                if f'Рисунок 4.{j+1}' in text or f'рисунок 4.{j+1}' in text.lower():
                    add_graph(result_doc, graphs_ch4 / fname, caption)
                    graph_idx = max(graph_idx, j + 1)
    
    # Добавляем оставшиеся графики главы 4
    print("    Добавление оставшихся графиков главы 4...")
    for fname, caption in ch4_graphs[graph_idx:]:
        add_graph(result_doc, graphs_ch4 / fname, caption)
    
    # ЧАСТЬ 4: Заключение, Список источников, Приложения
    print("  Копирование: Заключение, Список источников, Приложения...")
    start_part4 = sections.get('conclusion', 320)
    for i, para in enumerate(main_doc.paragraphs[start_part4:]):
        copy_para(para, result_doc)
    
    # Сохраняем результат
    print(f"\nСохранение: {output_path}")
    result_doc.save(output_path)
    
    size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"\nГотово! Размер: {size_mb:.2f} MB")
    
    # Проверяем количество изображений
    import zipfile
    with zipfile.ZipFile(output_path, 'r') as z:
        images = [f for f in z.namelist() if 'media' in f and f.endswith('.png')]
        print(f"Количество изображений: {len(images)}")
    
    return output_path

if __name__ == '__main__':
    main()
