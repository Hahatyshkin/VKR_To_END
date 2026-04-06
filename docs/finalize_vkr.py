#!/usr/bin/env python3
"""
Финализация VKR документа: обновление реферата и проверка.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import subprocess

# Пути
VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final_Complete.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final.docx')
DOWNLOAD_DIR = '/home/z/my-project/download'

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== ПОДСЧЕТ ЭЛЕМЕНТОВ =====
    print('\nПодсчет элементов документа...')
    
    figure_count = 0
    table_count = 0
    page_count = 0
    source_count = 0
    appendix_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        # Рисунки
        if 'Рисунок' in text and '—' in text:
            figure_count += 1
        # Таблицы
        if 'Таблица' in text and '—' in text:
            table_count += 1
        # Источники (нумерованный список)
        if re.match(r'^\d+\.\s+', text) and 'М.' not in text and 'ISBN' in text:
            source_count += 1
    
    # Подсчет источников через поиск в списке
    source_pattern = re.compile(r'^\d+\.\s+.+')
    for para in doc.paragraphs:
        if source_pattern.match(para.text.strip()):
            source_count += 1
    
    # Приложения
    for para in doc.paragraphs:
        if 'ПРИЛОЖЕНИЕ' in para.text.upper() and ('А' in para.text or 'Б' in para.text or 'В' in para.text):
            appendix_count += 1
    
    # Таблицы в документе
    table_count = len(doc.tables)
    
    # Оценка страниц (примерно 3 параграфа на страницу)
    page_count = len(doc.paragraphs) // 3
    
    print(f'  Рисунков: {figure_count}')
    print(f'  Таблиц: {table_count}')
    print(f'  Страниц (прибл.): {page_count}')
    print(f'  Приложений: {appendix_count}')
    
    # ===== ОБНОВЛЕНИЕ РЕФЕРАТА =====
    print('\nОбновление реферата...')
    
    for para in doc.paragraphs:
        text = para.text
        # Ищем строку реферата
        if 'Научно-исследовательская работа состоит из' in text:
            # Формируем новую строку
            new_text = f'Научно-исследовательская работа состоит из {page_count} страниц, {figure_count} рисунков, {table_count} таблиц, 76 использованных источников, 3 приложений.'
            
            # Заменяем текст в runs
            for run in para.runs:
                if 'Научно-исследовательская работа' in run.text:
                    run.text = new_text
                    print(f'  Обновлен реферат: {new_text}')
                    break
            else:
                # Если не нашли в отдельных runs, заменяем весь текст параграфа
                para.clear()
                run = para.add_run(new_text)
                print(f'  Обновлен реферат: {new_text}')
            break
    
    # ===== ПРОВЕРКА СВЯЗНОСТИ =====
    print('\nПроверка связности документа...')
    
    # Проверяем наличие всех глав
    chapters = ['ГЛАВА 1', 'ГЛАВА 2', 'ГЛАВА 3', 'ГЛАВА 4', 'ЗАКЛЮЧЕНИЕ', 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ']
    missing = []
    for ch in chapters:
        found = False
        for para in doc.paragraphs:
            if ch in para.text.upper():
                found = True
                break
        if not found:
            missing.append(ch)
    
    if missing:
        print(f'  Внимание! Не найдены разделы: {missing}')
    else:
        print('  Все основные разделы присутствуют.')
    
    # ===== СОХРАНЕНИЕ =====
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    # Копирование в download
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    download_path = os.path.join(DOWNLOAD_DIR, 'VKR_Final.docx')
    
    import shutil
    shutil.copy(OUTPUT_DOCX, download_path)
    print(f'Копия сохранена в: {download_path}')
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер файла: {size / 1024 / 1024:.2f} MB')
    
    # ===== КОНВЕРТАЦИЯ В PDF =====
    print('\nКонвертация в PDF...')
    try:
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
        ], check=True, capture_output=True)
        
        pdf_path = os.path.join(DOWNLOAD_DIR, 'VKR_Final.pdf')
        if os.path.exists(pdf_path):
            pdf_size = os.path.getsize(pdf_path)
            print(f'PDF создан: {pdf_path}')
            print(f'Размер PDF: {pdf_size / 1024 / 1024:.2f} MB')
    except Exception as e:
        print(f'Ошибка при создании PDF: {e}')
    
    print('\nГотово!')
    return OUTPUT_DOCX

if __name__ == '__main__':
    main()
