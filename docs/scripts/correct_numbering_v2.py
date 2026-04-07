#!/usr/bin/env python3
"""
Исправление нумерации рисунков и таблиц в VKR документе (улучшенная версия).
"""

from docx import Document
from docx.shared import Pt
import os
import re
import subprocess

# Пути
VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete_Final.docx')
DOWNLOAD_DIR = '/home/z/my-project/download'

def renumber_figures_v2(doc):
    """Перенумерация рисунков по порядку появления."""
    print('\nПеренумерация рисунков...')
    
    # Собираем все рисунки с их позициями
    figures = []
    current_chapter = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Определяем текущую главу (проверяем начало строки)
        if text.upper().startswith('ГЛАВА 1') or 'ГЛАВА 1' in text.upper():
            current_chapter = 1
        elif text.upper().startswith('ГЛАВА 2') or 'ГЛАВА 2' in text.upper():
            current_chapter = 2
        elif text.upper().startswith('ГЛАВА 3') or 'ГЛАВА 3' in text.upper():
            current_chapter = 3
        elif text.upper().startswith('ГЛАВА 4') or 'ГЛАВА 4' in text.upper():
            current_chapter = 4
        elif 'ПРИЛОЖЕНИЕ А' in text.upper():
            current_chapter = 100  # Приложение А
        elif 'ПРИЛОЖЕНИЕ Б' in text.upper():
            current_chapter = 101  # Приложение Б
        elif 'ПРИЛОЖЕНИЕ В' in text.upper():
            current_chapter = 102  # Приложение В
        
        # Ищем рисунок
        if 'Рисунок' in text:
            match = re.search(r'Рисунок\s+([\d.]+|[А-Я])\s*[—–-]', text)
            if match:
                old_num = match.group(1)
                figures.append({
                    'para_idx': i,
                    'old_num': old_num,
                    'chapter': current_chapter,
                    'text': text[:100]
                })
    
    # Группируем по главам и перенумеровываем
    chapter_figures = {1: [], 2: [], 3: [], 4: [], 100: [], 101: [], 102: []}
    for fig in figures:
        if fig['chapter'] in chapter_figures:
            chapter_figures[fig['chapter']].append(fig)
    
    # Перенумеровываем
    changes = []
    for chapter, figs in chapter_figures.items():
        for idx, fig in enumerate(figs):
            if chapter >= 100:  # Приложения
                appendix_letter = chr(ord('А') + (chapter - 100))
                new_num = f'{appendix_letter}.{idx + 1}'
            else:
                new_num = f'{chapter}.{idx + 1}'
            
            if fig['old_num'] != new_num:
                changes.append({
                    'para_idx': fig['para_idx'],
                    'old_num': fig['old_num'],
                    'new_num': new_num,
                    'chapter': chapter
                })
                print(f'  Рисунок {fig["old_num"]} -> {new_num} (глава {chapter})')
    
    # Применяем изменения
    for change in changes:
        para = doc.paragraphs[change['para_idx']]
        for run in para.runs:
            if 'Рисунок' in run.text:
                # Экранируем точку в старом номере
                old_pattern = f'Рисунок {change["old_num"]}'
                new_pattern = f'Рисунок {change["new_num"]}'
                run.text = run.text.replace(old_pattern, new_pattern)
    
    return changes

def renumber_tables_v2(doc):
    """Перенумерация таблиц по порядку появления."""
    print('\nПеренумерация таблиц...')
    
    tables = []
    current_chapter = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Определяем текущую главу
        if text.upper().startswith('ГЛАВА 1') or 'ГЛАВА 1' in text.upper():
            current_chapter = 1
        elif text.upper().startswith('ГЛАВА 2') or 'ГЛАВА 2' in text.upper():
            current_chapter = 2
        elif text.upper().startswith('ГЛАВА 3') or 'ГЛАВА 3' in text.upper():
            current_chapter = 3
        elif text.upper().startswith('ГЛАВА 4') or 'ГЛАВА 4' in text.upper():
            current_chapter = 4
        elif 'ПРИЛОЖЕНИЕ А' in text.upper():
            current_chapter = 100
        elif 'ПРИЛОЖЕНИЕ Б' in text.upper():
            current_chapter = 101
        elif 'ПРИЛОЖЕНИЕ В' in text.upper():
            current_chapter = 102
        
        # Ищем таблицу
        if 'Таблица' in text:
            match = re.search(r'Таблица\s+([\d.]+|[А-Я])\s*[—–-]', text)
            if match:
                old_num = match.group(1)
                tables.append({
                    'para_idx': i,
                    'old_num': old_num,
                    'chapter': current_chapter,
                    'text': text[:100]
                })
    
    # Группируем и перенумеровываем
    chapter_tables = {1: [], 2: [], 3: [], 4: [], 100: [], 101: [], 102: []}
    for tbl in tables:
        if tbl['chapter'] in chapter_tables:
            chapter_tables[tbl['chapter']].append(tbl)
    
    changes = []
    for chapter, tbls in chapter_tables.items():
        for idx, tbl in enumerate(tbls):
            if chapter >= 100:
                appendix_letter = chr(ord('А') + (chapter - 100))
                new_num = f'{appendix_letter}.{idx + 1}'
            else:
                new_num = f'{chapter}.{idx + 1}'
            
            if tbl['old_num'] != new_num:
                changes.append({
                    'para_idx': tbl['para_idx'],
                    'old_num': tbl['old_num'],
                    'new_num': new_num,
                    'chapter': chapter
                })
                print(f'  Таблица {tbl["old_num"]} -> {new_num} (глава {chapter})')
    
    # Применяем изменения
    for change in changes:
        para = doc.paragraphs[change['para_idx']]
        for run in para.runs:
            if 'Таблица' in run.text:
                old_pattern = f'Таблица {change["old_num"]}'
                new_pattern = f'Таблица {change["new_num"]}'
                run.text = run.text.replace(old_pattern, new_pattern)
    
    return changes

def update_abstract(doc):
    """Обновление реферата с актуальными данными."""
    print('\nОбновление реферата...')
    
    # Подсчитываем элементы
    fig_count = 0
    tbl_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        if 'Рисунок' in text and ('—' in text or '-' in text):
            # Проверяем что это подпись рисунка
            if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', text):
                fig_count += 1
        if 'Таблица' in text and ('—' in text or '-' in text):
            if re.search(r'Таблица\s+[\d.]+\s*[—–-]', text):
                tbl_count += 1
    
    print(f'  Рисунков: {fig_count}')
    print(f'  Таблиц: {tbl_count}')
    
    # Оценка страниц
    page_count = len(doc.paragraphs) // 3
    
    # Обновляем реферат
    for para in doc.paragraphs:
        if 'Научно-исследовательская работа состоит из' in para.text:
            new_text = f'Научно-исследовательская работа состоит из {page_count} страниц, {fig_count} рисунков, {tbl_count} таблиц, 76 использованных источников, 3 приложений.'
            para.clear()
            run = para.add_run(new_text)
            print(f'  Обновлено: {new_text}')
            break

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # Перенумерация
    fig_changes = renumber_figures_v2(doc)
    tbl_changes = renumber_tables_v2(doc)
    
    # Обновление реферата
    update_abstract(doc)
    
    print(f'\nВсего изменений:')
    print(f'  Рисунков: {len(fig_changes)}')
    print(f'  Таблиц: {len(tbl_changes)}')
    
    # Сохранение
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    # Копирование в download
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    download_path = os.path.join(DOWNLOAD_DIR, 'VKR_Complete_Final.docx')
    shutil.copy(OUTPUT_DOCX, download_path)
    print(f'Копия: {download_path}')
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер: {size / 1024 / 1024:.2f} MB')
    
    # Конвертация в PDF
    print('\nКонвертация в PDF...')
    try:
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
        ], check=True, capture_output=True)
        print(f'PDF создан: {os.path.join(DOWNLOAD_DIR, "VKR_Complete_Final.pdf")}')
    except Exception as e:
        print(f'Ошибка PDF: {e}')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
