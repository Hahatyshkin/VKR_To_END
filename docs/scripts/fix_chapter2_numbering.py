#!/usr/bin/env python3
"""
Финальное исправление нумерации рисунков в Главе 2.
"""

from docx import Document
from docx.shared import Pt
import os
import re
import subprocess

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_QUALITY_FINAL.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_FINAL_COMPLETE.docx')
DOWNLOAD_DIR = '/home/z/my-project/download'

def main():
    print('Исправление нумерации рисунков в Главе 2...\n')
    doc = Document(INPUT_DOCX)
    
    # Собираем все рисунки Главы 2 с их текущими номерами и позициями
    ch2_figs = []
    in_ch2 = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'ГЛАВА 2' in text.upper():
            in_ch2 = True
        elif 'ГЛАВА 3' in text.upper():
            in_ch2 = False
        
        if in_ch2 and 'Рисунок' in text:
            match = re.search(r'Рисунок\s+([\d.]+[а-я]?)', text)
            if match:
                old_num = match.group(1)
                ch2_figs.append({
                    'idx': i,
                    'old_num': old_num,
                    'para': para,
                    'text': text[:80]
                })
    
    print(f'Найдено {len(ch2_figs)} рисунков в Главе 2:\n')
    for fig in ch2_figs:
        print(f'  Рисунок {fig["old_num"]}: {fig["text"][:50]}...')
    
    # Перенумеровываем в правильном порядке
    print('\nИсправление нумерации:')
    for idx, fig in enumerate(ch2_figs):
        new_num = f'2.{idx + 1}'
        old_num = fig['old_num']
        
        if old_num != new_num:
            for run in fig['para'].runs:
                text = run.text
                if f'Рисунок {old_num}' in text:
                    run.text = text.replace(f'Рисунок {old_num}', f'Рисунок {new_num}')
                    print(f'  {old_num} -> {new_num}')
                    break
    
    # Сохранение
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    final_path = os.path.join(DOWNLOAD_DIR, 'VKR_Финал.docx')
    shutil.copy(OUTPUT_DOCX, final_path)
    
    # PDF
    print('Конвертация в PDF...')
    subprocess.run([
        'soffice', '--headless', '--convert-to', 'pdf',
        '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
    ], check=True, capture_output=True)
    
    # Проверка результата
    print('\n=== ПРОВЕРКА РЕЗУЛЬТАТА ===')
    doc2 = Document(OUTPUT_DOCX)
    
    # Подсчет
    total_figs = 0
    total_tables = 0
    
    for para in doc2.paragraphs:
        if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', para.text):
            total_figs += 1
        if re.search(r'Таблица\s+[\d.]+\s*[—–-]', para.text):
            total_tables += 1
    
    print(f'Рисунков: {total_figs}')
    print(f'Таблиц: {total_tables}')
    
    # Проверка Главы 2
    print('\nРисунки Главы 2:')
    in_ch2 = False
    for para in doc2.paragraphs:
        text = para.text.strip()
        if 'ГЛАВА 2' in text.upper():
            in_ch2 = True
        elif 'ГЛАВА 3' in text.upper():
            in_ch2 = False
        
        if in_ch2 and 'Рисунок 2.' in text:
            match = re.search(r'Рисунок\s+2\.\d+\s*[—–-]', text)
            if match:
                print(f'  {match.group()}')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
