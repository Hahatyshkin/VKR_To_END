#!/usr/bin/env python3
"""
Комплексная проверка и исправление VKR документа.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import subprocess
from copy import deepcopy

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_QUALITY_FINAL.docx')
NEW_FIGURES_DIR = os.path.join(VKR_DIR, 'media/new_figures')
DOWNLOAD_DIR = '/home/z/my-project/download'

def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_table_simple(doc, headers, rows, caption_text, insert_after_para):
    """Добавление таблицы после указанного параграфа."""
    # Создаем таблицу в конце документа
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(11)
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    # Границы
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    for row_idx, row in enumerate(rows):
        for col_idx, cell_data in enumerate(row):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    space = doc.add_paragraph()
    
    # Перемещаем после указанного параграфа
    insert_after_para._element.addnext(space._element)
    insert_after_para._element.addnext(table._tbl)
    insert_after_para._element.addnext(caption._element)
    
    return table

def add_figure_simple(doc, image_path, caption_text, insert_after_para):
    """Добавление рисунка после указанного параграфа."""
    # Создаем параграфы в конце
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(image_path, width=Inches(5.5))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(11)
    
    space = doc.add_paragraph()
    
    # Перемещаем
    insert_after_para._element.addnext(space._element)
    insert_after_para._element.addnext(caption._element)
    insert_after_para._element.addnext(fig_para._element)

def main():
    print('='*70)
    print('КОМПЛЕКСНАЯ ПРОВЕРКА И ИСПРАВЛЕНИЕ VKR ДОКУМЕНТА')
    print('='*70)
    
    print('\nОткрытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== АНАЛИЗ ТЕКУЩЕГО СОСТОЯНИЯ =====
    print('\n--- АНАЛИЗ ТЕКУЩЕГО СОСТОЯНИЯ ---')
    
    # Подсчет рисунков по главам
    chapters_figs = {1: 0, 2: 0, 3: 0, 4: 0}
    current_chapter = 0
    fig_list = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'ГЛАВА 1' in text.upper():
            current_chapter = 1
        elif 'ГЛАВА 2' in text.upper():
            current_chapter = 2
        elif 'ГЛАВА 3' in text.upper():
            current_chapter = 3
        elif 'ГЛАВА 4' in text.upper():
            current_chapter = 4
        elif 'ПРИЛОЖЕНИЕ' in text.upper():
            current_chapter = 10
        
        if 'Рисунок' in text and current_chapter in chapters_figs:
            match = re.search(r'Рисунок\s+([\d.]+)', text)
            if match:
                chapters_figs[current_chapter] += 1
                fig_list.append((current_chapter, match.group(1), text[:60]))
    
    print(f'\nРисунки по главам:')
    for ch, count in chapters_figs.items():
        print(f'  Глава {ch}: {count} рисунков')
    
    # Подсчет таблиц
    chapters_tables = {1: 0, 2: 0, 3: 0, 4: 0}
    current_chapter = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if 'ГЛАВА 1' in text.upper():
            current_chapter = 1
        elif 'ГЛАВА 2' in text.upper():
            current_chapter = 2
        elif 'ГЛАВА 3' in text.upper():
            current_chapter = 3
        elif 'ГЛАВА 4' in text.upper():
            current_chapter = 4
        
        if 'Таблица' in text and current_chapter in chapters_tables:
            match = re.search(r'Таблица\s+[\d.]+', text)
            if match:
                chapters_tables[current_chapter] += 1
    
    print(f'\nТаблицы по главам:')
    for ch, count in chapters_tables.items():
        print(f'  Глава {ch}: {count} таблиц')
    
    # ===== ИСПРАВЛЕНИЕ НУМЕРАЦИИ В ГЛАВЕ 2 =====
    print('\n--- ИСПРАВЛЕНИЕ НУМЕРАЦИИ В ГЛАВЕ 2 ---')
    
    # Сначала находим все рисунки в Главе 2 и исправляем
    ch2_figures = []
    in_ch2 = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if 'ГЛАВА 2' in text.upper():
            in_ch2 = True
        elif 'ГЛАВА 3' in text.upper():
            in_ch2 = False
        
        if in_ch2 and 'Рисунок' in text:
            ch2_figures.append((i, para))
    
    # Заменяем нумерацию на последовательную
    for idx, (para_idx, para) in enumerate(ch2_figures):
        new_num = f'2.{idx + 1}'
        for run in para.runs:
            text = run.text
            # Ищем старый номер рисунка
            match = re.search(r'Рисунок\s+([\d.]+)', text)
            if match:
                old_num = match.group(1)
                if old_num != new_num:
                    run.text = text.replace(f'Рисунок {old_num}', f'Рисунок {new_num}')
                    print(f'  Рисунок {old_num} -> {new_num}')
    
    # ===== ПОИСК ПОЗИЦИЙ ДЛЯ ВСТАВКИ =====
    print('\n--- ПОИСК ПОЗИЦИЙ ДЛЯ ВСТАВКИ МАТЕРИАЛОВ ---')
    
    positions = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Конец раздела 1.3.6 (перед 1.4)
        if '1.4' in text and 'Интеграция' in text:
            positions['before_14'] = i
            print(f'  Перед разделом 1.4: параграф {i}')
        
        # После описания FWHT (для матриц Адамара)
        if 'H₄ = [1 1 1 1' in text or 'H4 = [1 1 1 1' in text or ('рекурсивно' in text.lower() and 'адамара' in text.lower()):
            positions['after_hadamard_desc'] = i
            print(f'  После описания матриц Адамара: параграф {i}')
        
        # После описания преимуществ FWHT
        if 'ограниченная применимость' in text.lower() and 'спектральн' in text.lower():
            positions['after_fwht_pros'] = i
            print(f'  После преимуществ FWHT: параграф {i}')
        
        # Перед 1.5.4 (для таблицы метрик)
        if '1.5.4' in text and 'Характеристики' in text:
            positions['before_154'] = i
            print(f'  Перед разделом 1.5.4: параграф {i}')
        
        # Перед выводами главы 1
        if '1.6' in text and 'Выводы' in text:
            positions['before_16'] = i
            print(f'  Перед разделом 1.6: параграф {i}')
    
    # ===== ВСТАВКА ТАБЛИЦ =====
    print('\n--- ВСТАВКА ТАБЛИЦ В ГЛАВУ 1 ---')
    
    # Таблица 1.1: Вычислительная сложность
    headers1 = ['Метод', 'Операции', 'Сложность', 'Умножения']
    rows1 = [
        ['FFT', 'N·log₂(N)', 'O(N log N)', 'Комплексные'],
        ['FWHT', 'N·log₂(N)', 'O(N log N)', 'Нет'],
        ['DCT', 'N·log₂(N)', 'O(N log N)', 'Вещественные'],
        ['DWT (Хаар)', 'N', 'O(N)', 'Нет'],
        ['μ-law', 'N', 'O(N)', 'Нет'],
    ]
    
    if 'before_14' in positions:
        try:
            para = doc.paragraphs[positions['before_14'] - 1]
            add_table_simple(doc, headers1, rows1, 
                'Таблица 1.1 — Сравнение вычислительной сложности методов преобразования',
                para)
            print('  ✓ Таблица 1.1 добавлена')
        except Exception as e:
            print(f'  ✗ Ошибка таблицы 1.1: {e}')
    
    # Таблица 1.2: Метрики качества
    headers2 = ['Метрика', 'Область', 'Диапазон', 'Описание']
    rows2 = [
        ['SNR', 'Временная', '0–∞ дБ', 'Отношение сигнал/шум'],
        ['SI-SDR', 'Временная', '–∞–∞ дБ', 'Масштабно-инвариантное'],
        ['RMSE', 'Временная', '0–∞', 'Среднеквадратичная ошибка'],
        ['LSD', 'Спектральная', '0–∞ дБ', 'Спектральное расстояние'],
        ['STOI', 'Психоакуст.', '0–1', 'Разборчивость речи'],
        ['PESQ', 'Психоакуст.', '–0.5–4.5', 'Качество речи'],
    ]
    
    if 'before_154' in positions:
        try:
            para = doc.paragraphs[positions['before_154'] - 1]
            add_table_simple(doc, headers2, rows2,
                'Таблица 1.2 — Характеристики метрик качества аудиосигналов',
                para)
            print('  ✓ Таблица 1.2 добавлена')
        except Exception as e:
            print(f'  ✗ Ошибка таблицы 1.2: {e}')
    
    # ===== ВСТАВКА РИСУНКОВ =====
    print('\n--- ВСТАВКА РИСУНКОВ В ГЛАВУ 1 ---')
    
    # Рисунок 1.4: Матрицы Адамара
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_hadamard_matrices.png')
    if os.path.exists(fig_path) and 'after_hadamard_desc' in positions:
        try:
            para = doc.paragraphs[positions['after_hadamard_desc']]
            add_figure_simple(doc, fig_path, 
                'Рисунок 1.4 — Матрицы Адамара порядков 1, 2, 4, 8', para)
            print('  ✓ Рисунок 1.4 добавлен')
        except Exception as e:
            print(f'  ✗ Ошибка рисунка 1.4: {e}')
    
    # Рисунок 1.5: Функции Уолша
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_walsh_functions.png')
    if os.path.exists(fig_path) and 'after_fwht_pros' in positions:
        try:
            para = doc.paragraphs[positions['after_fwht_pros']]
            add_figure_simple(doc, fig_path,
                'Рисунок 1.5 — Функции Уолша (первые 8 функций)', para)
            print('  ✓ Рисунок 1.5 добавлен')
        except Exception as e:
            print(f'  ✗ Ошибка рисунка 1.5: {e}')
    
    # Рисунок 1.6: Сравнение FFT/FWHT
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_fft_vs_fwht.png')
    if os.path.exists(fig_path) and 'before_14' in positions:
        try:
            para = doc.paragraphs[positions['before_14'] - 3]
            add_figure_simple(doc, fig_path,
                'Рисунок 1.6 — Сравнение спектров FFT и FWHT', para)
            print('  ✓ Рисунок 1.6 добавлен')
        except Exception as e:
            print(f'  ✗ Ошибка рисунка 1.6: {e}')
    
    # Рисунок 1.7: Схема обработки
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_audio_pipeline.png')
    if os.path.exists(fig_path) and 'before_16' in positions:
        try:
            para = doc.paragraphs[positions['before_16'] - 1]
            add_figure_simple(doc, fig_path,
                'Рисунок 1.7 — Общая схема обработки аудиосигнала', para)
            print('  ✓ Рисунок 1.7 добавлен')
        except Exception as e:
            print(f'  ✗ Ошибка рисунка 1.7: {e}')
    
    # ===== ФИНАЛЬНЫЙ ПОДСЧЕТ =====
    print('\n--- ФИНАЛЬНЫЙ ПОДСЧЕТ ---')
    
    total_figs = 0
    total_tables = 0
    
    for para in doc.paragraphs:
        if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', para.text):
            total_figs += 1
        if re.search(r'Таблица\s+[\d.]+\s*[—–-]', para.text):
            total_tables += 1
    
    page_count = len(doc.paragraphs) // 3
    
    print(f'  Страниц: ~{page_count}')
    print(f'  Рисунков: {total_figs}')
    print(f'  Таблиц: {total_tables}')
    print(f'  Источников: 76')
    print(f'  Приложений: 3')
    
    # Обновление реферата
    for para in doc.paragraphs:
        if 'Научно-исследовательская работа состоит из' in para.text:
            para.clear()
            para.add_run(f'Научно-исследовательская работа состоит из {page_count} страниц, {total_figs} рисунков, {total_tables} таблиц, 76 использованных источников, 3 приложений.')
            print('  Реферат обновлен')
            break
    
    # ===== СОХРАНЕНИЕ =====
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    final_path = os.path.join(DOWNLOAD_DIR, 'VKR_Качественный_Финал.docx')
    shutil.copy(OUTPUT_DOCX, final_path)
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер: {size / 1024 / 1024:.2f} MB')
    
    # PDF
    print('\nКонвертация в PDF...')
    subprocess.run([
        'soffice', '--headless', '--convert-to', 'pdf',
        '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
    ], check=True, capture_output=True)
    
    print('\n' + '='*70)
    print('ГОТОВО!')
    print('='*70)

if __name__ == '__main__':
    main()
