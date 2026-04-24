#!/usr/bin/env python3
"""
Создание документа с дополнительными материалами для ВКР.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = '/home/z/my-project/download'
NEW_FIGURES_DIR = '/home/z/my-project/VKR/docs/media/new_figures'

def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, caption_text):
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(11)
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    for row_idx, row in enumerate(rows):
        for col_idx, cell_data in enumerate(row):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()

def add_figure(doc, image_path, caption_text, width=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()

def main():
    print('Создание документа с дополнительными материалами...')
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Дополнительные материалы для ВКР', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    intro = doc.add_paragraph()
    intro.add_run('Данный документ содержит рисунки и таблицы, которые рекомендуется вставить в соответствующие разделы Главы 1 ВКР.')
    
    # === РИСУНКИ ===
    doc.add_heading('Рисунки для Главы 1', 1)
    
    # Рисунок 1.4
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_hadamard_matrices.png')
    if os.path.exists(fig_path):
        add_figure(doc, fig_path, 
            'Рисунок 1.4 — Матрицы Адамара порядков 1, 2, 4, 8')
        doc.add_paragraph('Вставить после раздела 1.3.2 (после описания рекурсивного построения матриц Адамара).')
    
    # Рисунок 1.5
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_walsh_functions.png')
    if os.path.exists(fig_path):
        add_figure(doc, fig_path,
            'Рисунок 1.5 — Функции Уолша (первые 8 функций)')
        doc.add_paragraph('Вставить после раздела 1.3.2 (после описания преимуществ FWHT).')
    
    # Рисунок 1.6
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_fft_vs_fwht.png')
    if os.path.exists(fig_path):
        add_figure(doc, fig_path,
            'Рисунок 1.6 — Сравнение спектров FFT и FWHT')
        doc.add_paragraph('Вставить после описания различий между FFT и FWHT.')
    
    # Рисунок 1.7
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_audio_pipeline.png')
    if os.path.exists(fig_path):
        add_figure(doc, fig_path,
            'Рисунок 1.7 — Общая схема обработки аудиосигнала')
        doc.add_paragraph('Вставить в раздел 1.4 (Интеграция методов преобразования).')
    
    # Рисунок 1.8
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_energy_concentration.png')
    if os.path.exists(fig_path):
        add_figure(doc, fig_path,
            'Рисунок 1.8 — Концентрация энергии в коэффициентах преобразований')
        doc.add_paragraph('Вставить после описания методов отбора коэффициентов.')
    
    # === ТАБЛИЦЫ ===
    doc.add_heading('Таблицы для Главы 1', 1)
    
    # Таблица 1.1
    doc.add_heading('Таблица 1.1 — Сравнение вычислительной сложности методов', 2)
    headers1 = ['Метод', 'Операции', 'Сложность', 'Умножения']
    rows1 = [
        ['FFT', 'N·log₂(N)', 'O(N log N)', 'Комплексные'],
        ['FWHT', 'N·log₂(N)', 'O(N log N)', 'Нет'],
        ['DCT', 'N·log₂(N)', 'O(N log N)', 'Вещественные'],
        ['DWT (Хаар)', 'N', 'O(N)', 'Нет'],
        ['μ-law', 'N', 'O(N)', 'Нет'],
    ]
    add_table(doc, headers1, rows1, 'Таблица 1.1 — Сравнение вычислительной сложности методов преобразования')
    doc.add_paragraph('Вставить после раздела 1.3.6 (перед разделом 1.4).')
    
    # Таблица 1.2
    doc.add_heading('Таблица 1.2 — Метрики качества аудиосигналов', 2)
    headers2 = ['Метрика', 'Область', 'Диапазон', 'Описание']
    rows2 = [
        ['SNR', 'Временная', '0–∞ дБ', 'Отношение сигнал/шум'],
        ['SI-SDR', 'Временная', '–∞–∞ дБ', 'Масштабно-инвариантное отношение сигнал/искажение'],
        ['RMSE', 'Временная', '0–∞', 'Среднеквадратичная ошибка'],
        ['LSD', 'Спектральная', '0–∞ дБ', 'Логарифмическое спектральное расстояние'],
        ['Spectral Conv.', 'Спектральная', '0–1', 'Близость амплитудных спектров'],
        ['STOI', 'Психоакуст.', '0–1', 'Индекс разборчивости речи'],
        ['PESQ', 'Психоакуст.', '–0.5–4.5', 'Оценка качества речи (ITU-T P.862)'],
    ]
    add_table(doc, headers2, rows2, 'Таблица 1.2 — Характеристики метрик качества аудиосигналов')
    doc.add_paragraph('Вставить перед разделом 1.5.4 (Характеристики сигнала).')
    
    # Таблица 1.3
    doc.add_heading('Таблица 1.3 — Сравнительный анализ методов преобразования', 2)
    headers3 = ['Метод', 'Преимущества', 'Недостатки', 'Применение']
    rows3 = [
        ['FFT', 'Точный спектральный анализ, стандарт IEEE', 'Комплексные вычисления, требует умножений', 'Спектральный анализ, фильтрация'],
        ['FWHT', 'Быстрый алгоритм, нет умножений, простая реализация', 'Ограниченная точность для сложных сигналов', 'Сжатие, кодирование, криптография'],
        ['DCT', 'Концентрация энергии, стандарт в JPEG/MP3', 'Требует вещественных умножений', 'Сжатие изображений и аудио'],
        ['DWT', 'Локализация по времени и частоте, многомасштабность', 'Сложная реализация, выбор базиса', 'Анализ переходных процессов'],
        ['μ-law', 'Простота, логарифмическая шкала, улучшение SNR', 'Узкая область применения', 'Телефония, телекоммуникации'],
    ]
    add_table(doc, headers3, rows3, 'Таблица 1.3 — Сравнительный анализ методов преобразования сигналов')
    doc.add_paragraph('Вставить в раздел 1.3 (Методы преобразования сигналов для сжатия).')
    
    # Сохранение
    output_path = os.path.join(OUTPUT_DIR, 'VKR_Дополнительные_материалы.docx')
    doc.save(output_path)
    print(f'Документ сохранен: {output_path}')
    print(f'Размер: {os.path.getsize(output_path) / 1024:.1f} KB')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
