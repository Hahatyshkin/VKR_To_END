#!/usr/bin/env python3
"""
Исправление нумерации рисунков и таблиц в VKR документе.
"""

from docx import Document
from docx.shared import Pt
import os
import re

# Пути
VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Corrected.docx')
DOWNLOAD_DIR = '/home/z/my-project/download'

def renumber_figures(doc):
    """Перенумерация рисунков по порядку появления в каждой главе."""
    print('\nПеренумерация рисунков...')
    
    # Собираем все рисунки с их позициями
    figures = []
    current_chapter = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Определяем текущую главу
        if text.startswith('ГЛАВА 1'):
            current_chapter = 1
        elif text.startswith('ГЛАВА 2'):
            current_chapter = 2
        elif text.startswith('ГЛАВА 3'):
            current_chapter = 3
        elif text.startswith('ГЛАВА 4'):
            current_chapter = 4
        elif text.startswith('ПРИЛОЖЕНИЕ'):
            current_chapter = 10  # Приложения
        
        # Ищем рисунок
        if 'Рисунок' in text and '—' in text:
            match = re.search(r'Рисунок\s+([\d.]+)\s*—', text)
            if match:
                old_num = match.group(1)
                figures.append({
                    'para_idx': i,
                    'old_num': old_num,
                    'chapter': current_chapter,
                    'text': text
                })
    
    # Группируем по главам и перенумеровываем
    chapter_figures = {1: [], 2: [], 3: [], 4: [], 10: []}
    for fig in figures:
        if fig['chapter'] in chapter_figures:
            chapter_figures[fig['chapter']].append(fig)
    
    # Перенумеровываем
    changes = []
    for chapter, figs in chapter_figures.items():
        for idx, fig in enumerate(figs):
            if chapter == 10:  # Приложения
                new_num = chr(ord('А') + idx)  # А, Б, В...
            else:
                new_num = f'{chapter}.{idx + 1}'
            
            if fig['old_num'] != new_num:
                changes.append({
                    'para_idx': fig['para_idx'],
                    'old_num': fig['old_num'],
                    'new_num': new_num
                })
                print(f'  Рисунок {fig["old_num"]} -> {new_num}')
    
    # Применяем изменения
    for change in changes:
        para = doc.paragraphs[change['para_idx']]
        for run in para.runs:
            if 'Рисунок' in run.text:
                old_pattern = f'Рисунок {change["old_num"]}'
                new_pattern = f'Рисунок {change["new_num"]}'
                run.text = run.text.replace(old_pattern, new_pattern)
    
    return len(changes)

def renumber_tables(doc):
    """Перенумерация таблиц по порядку появления в каждой главе."""
    print('\nПеренумерация таблиц...')
    
    tables = []
    current_chapter = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Определяем текущую главу
        if text.startswith('ГЛАВА 1'):
            current_chapter = 1
        elif text.startswith('ГЛАВА 2'):
            current_chapter = 2
        elif text.startswith('ГЛАВА 3'):
            current_chapter = 3
        elif text.startswith('ГЛАВА 4'):
            current_chapter = 4
        elif text.startswith('ПРИЛОЖЕНИЕ'):
            current_chapter = 10
        
        # Ищем таблицу
        if 'Таблица' in text and '—' in text:
            match = re.search(r'Таблица\s+([\d.]+)\s*—', text)
            if match:
                old_num = match.group(1)
                tables.append({
                    'para_idx': i,
                    'old_num': old_num,
                    'chapter': current_chapter,
                    'text': text
                })
    
    # Группируем и перенумеровываем
    chapter_tables = {1: [], 2: [], 3: [], 4: [], 10: []}
    for tbl in tables:
        if tbl['chapter'] in chapter_tables:
            chapter_tables[tbl['chapter']].append(tbl)
    
    changes = []
    for chapter, tbls in chapter_tables.items():
        for idx, tbl in enumerate(tbls):
            if chapter == 10:
                new_num = chr(ord('А') + idx)
            else:
                new_num = f'{chapter}.{idx + 1}'
            
            if tbl['old_num'] != new_num:
                changes.append({
                    'para_idx': tbl['para_idx'],
                    'old_num': tbl['old_num'],
                    'new_num': new_num
                })
                print(f'  Таблица {tbl["old_num"]} -> {new_num}')
    
    # Применяем изменения
    for change in changes:
        para = doc.paragraphs[change['para_idx']]
        for run in para.runs:
            if 'Таблица' in run.text:
                old_pattern = f'Таблица {change["old_num"]}'
                new_pattern = f'Таблица {change["new_num"]}'
                run.text = run.text.replace(old_pattern, new_pattern)
    
    return len(changes)

def update_references(doc, fig_changes, tbl_changes):
    """Обновление ссылок на рисунки и таблицы в тексте."""
    print('\nОбновление ссылок в тексте...')
    
    # Объединяем все изменения
    all_changes = {}
    for c in fig_changes:
        all_changes[f'рисунок {c["old_num"]}'] = f'рисунок {c["new_num"]}'
        all_changes[f'Рисунок {c["old_num"]}'] = f'Рисунок {c["new_num"]}'
    for c in tbl_changes:
        all_changes[f'таблица {c["old_num"]}'] = f'таблица {c["new_num"]}'
        all_changes[f'Таблица {c["old_num"]}'] = f'Таблица {c["new_num"]}'
    
    # Ищем и обновляем ссылки
    refs_updated = 0
    for para in doc.paragraphs:
        for run in para.runs:
            original = run.text
            updated = original
            for old, new in all_changes.items():
                if old.lower() in updated.lower():
                    updated = re.sub(re.escape(old), new, updated, flags=re.IGNORECASE)
            if updated != original:
                run.text = updated
                refs_updated += 1
    
    print(f'  Обновлено ссылок: {refs_updated}')

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # Перенумерация
    fig_changes = renumber_figures(doc)
    tbl_changes = renumber_tables(doc)
    
    print(f'\nВсего изменений:')
    print(f'  Рисунков: {fig_changes}')
    print(f'  Таблиц: {tbl_changes}')
    
    # Сохранение
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    # Копирование в download
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    download_path = os.path.join(DOWNLOAD_DIR, 'VKR_Corrected.docx')
    shutil.copy(OUTPUT_DOCX, download_path)
    print(f'Копия: {download_path}')
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер: {size / 1024 / 1024:.2f} MB')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
