#!/usr/bin/env python3
"""
Скрипт для объединения ВКР с расширенными главами и встроенными изображениями.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import shutil
from pathlib import Path

def set_cell_margins(cell, top=100, bottom=100, left=180, right=180):
    """Установка отступов в ячейке таблицы"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_image_to_paragraph(paragraph, image_path, width_inches=5.5):
    """Добавление изображения в параграф"""
    if os.path.exists(image_path):
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        return True
    return False

def create_caption(paragraph, text, is_italic=True):
    """Создание подписи к рисунку/таблице"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    if is_italic:
        run.italic = True
    run.font.size = Pt(10)

def main():
    base_dir = Path('/home/z/my-project/VKR')
    docs_dir = base_dir / 'docs'
    final_dir = docs_dir / 'final'
    output_path = docs_dir / 'VKR_Complete.docx'
    
    # Пути к графикам
    graphs_ch3 = final_dir / 'graphs_chapter3'
    graphs_ch4 = final_dir / 'graphs_chapter4'
    
    print("Загрузка основного документа VKR.docx...")
    main_doc = Document(docs_dir / 'VKR.docx')
    
    print("Загрузка расширенной главы 3...")
    ch3_doc = Document(final_dir / 'VKR_Chapter3_Extended.docx')
    
    print("Загрузка расширенной главы 4...")
    ch4_doc = Document(final_dir / 'VKR_Chapter4_Extended.docx')
    
    # Поиск начала и конца глав в основном документе
    ch3_start = None
    ch3_end = None
    ch4_start = None
    ch4_end = None
    
    for i, para in enumerate(main_doc.paragraphs):
        text = para.text.strip()
        if text.startswith('ГЛАВА 3') or text.startswith('[ГЛАВА 3'):
            ch3_start = i
        elif text.startswith('ГЛАВА 4') or text.startswith('[ГЛАВА 4'):
            ch3_end = i
            ch4_start = i
        elif text.startswith('ЗАКЛЮЧЕНИЕ') or text.startswith('[ЗАКЛЮЧЕНИЕ'):
            ch4_end = i
    
    print(f"Глава 3: параграфы {ch3_start} - {ch3_end}")
    print(f"Глава 4: параграфы {ch4_start} - {ch4_end}")
    
    # Создаём новый документ на основе основного
    print("Создание объединённого документа...")
    
    # Копируем оригинальный документ
    shutil.copy(docs_dir / 'VKR.docx', output_path)
    
    # Открываем копию
    result_doc = Document(output_path)
    
    # Находим позиции глав в результирующем документе
    ch3_start_idx = None
    ch4_start_idx = None
    conclusion_idx = None
    
    for i, para in enumerate(result_doc.paragraphs):
        text = para.text.strip()
        if text.startswith('ГЛАВА 3') or 'ГЛАВА 3' in text:
            ch3_start_idx = i
        elif text.startswith('ГЛАВА 4') or 'ГЛАВА 4' in text:
            ch4_start_idx = i
        elif text.startswith('ЗАКЛЮЧЕНИЕ') or 'ЗАКЛЮЧЕНИЕ' in text:
            conclusion_idx = i
    
    print(f"Позиции в новом документе: Г3={ch3_start_idx}, Г4={ch4_start_idx}, Закл={conclusion_idx}")
    
    # Вставляем изображения в главу 3
    # Карта изображений для главы 3
    ch3_images = {
        'Рисунок 3.1': graphs_ch3 / 'snr_comparison.png',
        'Рисунок 3.2': graphs_ch3 / 'sisdr_comparison.png',
        'Рисунок 3.3': graphs_ch3 / 'lsd_comparison.png',
        'Рисунок 3.4': graphs_ch3 / 'time_comparison.png',
        'Рисунок 3.5': graphs_ch3 / 'size_comparison.png',
        'Рисунок 3.6': graphs_ch3 / 'snr_vs_time.png',
        'Рисунок 3.7': graphs_ch3 / 'snr_vs_size.png',
        'Рисунок 3.8': graphs_ch3 / 'radar_comparison.png',
        'Рисунок 3.9': graphs_ch3 / 'snr_boxplot.png',
        'Рисунок 3.10': graphs_ch3 / 'correlation_heatmap.png',
    }
    
    # Карта изображений для главы 4
    ch4_images = {
        'Рисунок 4.1': graphs_ch4 / 'snr_comparison_all.png',
        'Рисунок 4.2': graphs_ch4 / 'time_comparison_all.png',
        'Рисунок 4.3': graphs_ch4 / 'snr_vs_time.png',
        'Рисунок 4.4': graphs_ch4 / 'lsd_comparison.png',
        'Рисунок 4.5': graphs_ch4 / 'radar_comparison.png',
    }
    
    # Ищем параграфы с упоминаниями рисунков и вставляем изображения после них
    for i, para in enumerate(result_doc.paragraphs):
        text = para.text
        
        # Для главы 3
        for fig_name, fig_path in ch3_images.items():
            if fig_name in text and '![' not in text:
                # Ищем следующий параграф с "Рисунок" или создаём новый
                if fig_path.exists():
                    # Добавляем изображение после параграфа с таблицей или перед подписью
                    print(f"  Вставка {fig_name}: {fig_path}")
                    try:
                        # Создаём новый параграф с изображением
                        new_para = para.insert_paragraph_before('')
                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = new_para.add_run()
                        run.add_picture(str(fig_path), width=Inches(5.5))
                    except Exception as e:
                        print(f"    Ошибка: {e}")
        
        # Для главы 4
        for fig_name, fig_path in ch4_images.items():
            if fig_name in text and '![' not in text:
                if fig_path.exists():
                    print(f"  Вставка {fig_name}: {fig_path}")
                    try:
                        new_para = para.insert_paragraph_before('')
                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = new_para.add_run()
                        run.add_picture(str(fig_path), width=Inches(5.5))
                    except Exception as e:
                        print(f"    Ошибка: {e}")
    
    # Сохраняем результат
    print(f"Сохранение документа: {output_path}")
    result_doc.save(output_path)
    
    # Проверяем размер
    size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"Размер итогового файла: {size_mb:.2f} MB")
    
    print("Готово!")

if __name__ == '__main__':
    main()
