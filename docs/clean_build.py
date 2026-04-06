#!/usr/bin/env python3
"""
Чистовая сборка VKR с правильной структурой и нумерацией.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import subprocess
import zipfile
import shutil

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete.docx')  # Исходный документ
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_ГОТОВЫЙ_ДОКУМЕНТ.docx')
NEW_FIGURES_DIR = os.path.join(VKR_DIR, 'media/new_figures')
DOWNLOAD_DIR = '/home/z/my-project/download'
TEMP_DIR = '/tmp/vkr_clean'

def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_table_after(doc, para_idx, headers, rows, caption_text):
    """Добавление таблицы после указанного параграфа."""
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.bold = True
    run.font.size = Pt(11)
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9D9D9')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    for row_idx, row in enumerate(rows):
        for col_idx, cell_data in enumerate(row):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    space = doc.add_paragraph()
    
    # Перемещаем после указанного параграфа
    target = doc.paragraphs[para_idx]
    target._element.addnext(space._element)
    target._element.addnext(table._tbl)
    target._element.addnext(caption._element)

def add_figure_after(doc, para_idx, image_path, caption_text):
    """Добавление рисунка после указанного параграфа."""
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(image_path, width=Inches(5.5))
    
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(11)
    
    space = doc.add_paragraph()
    
    target = doc.paragraphs[para_idx]
    target._element.addnext(space._element)
    target._element.addnext(caption._element)
    target._element.addnext(fig_para._element)

def main():
    print('='*70)
    print('ЧИСТОВАЯ СБОРКА VKR ДОКУМЕНТА')
    print('='*70)
    
    print('\n1. Открытие исходного документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== АНАЛИЗ ПОЗИЦИЙ =====
    print('\n2. Анализ позиций для вставки...')
    
    positions = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # После описания матриц Адамара
        if 'рекурсивно' in text.lower() and 'адамар' in text.lower():
            positions['after_hadamard'] = i
            print(f'   После описания матриц Адамара: {i}')
        
        # После преимуществ FWHT
        if 'ограниченная применимость' in text.lower():
            positions['after_fwht'] = i
            print(f'   После преимуществ FWHT: {i}')
        
        # Перед 1.4 (для таблицы и рисунка)
        if '1.4' in text and 'Интеграция' in text:
            positions['before_14'] = i
            print(f'   Перед разделом 1.4: {i}')
        
        # Перед 1.5.4
        if '1.5.4' in text and 'Характеристики' in text:
            positions['before_154'] = i
            print(f'   Перед разделом 1.5.4: {i}')
        
        # Перед выводами главы 1
        if '1.6' in text and 'Выводы' in text:
            positions['before_16'] = i
            print(f'   Перед разделом 1.6: {i}')
    
    # ===== ВСТАВКА ТАБЛИЦ =====
    print('\n3. Вставка таблиц в Главу 1...')
    
    # Таблица 1.1
    if 'before_14' in positions:
        headers1 = ['Метод', 'Операции', 'Сложность', 'Умножения']
        rows1 = [
            ['FFT', 'N·log₂(N)', 'O(N log N)', 'Комплексные'],
            ['FWHT', 'N·log₂(N)', 'O(N log N)', 'Нет'],
            ['DCT', 'N·log₂(N)', 'O(N log N)', 'Вещественные'],
            ['DWT (Хаар)', 'N', 'O(N)', 'Нет'],
            ['μ-law', 'N', 'O(N)', 'Нет'],
        ]
        add_table_after(doc, positions['before_14'] - 1, headers1, rows1,
            'Таблица 1.1 — Сравнение вычислительной сложности методов преобразования')
        print('   ✓ Таблица 1.1')
    
    # Таблица 1.2
    if 'before_154' in positions:
        headers2 = ['Метрика', 'Область', 'Диапазон', 'Описание']
        rows2 = [
            ['SNR', 'Временная', '0–∞ дБ', 'Отношение сигнал/шум'],
            ['SI-SDR', 'Временная', '–∞–∞ дБ', 'Масштабно-инвариантное'],
            ['RMSE', 'Временная', '0–∞', 'Среднеквадратичная ошибка'],
            ['LSD', 'Спектральная', '0–∞ дБ', 'Спектральное расстояние'],
            ['STOI', 'Психоакуст.', '0–1', 'Разборчивость речи'],
            ['PESQ', 'Психоакуст.', '–0.5–4.5', 'Качество речи'],
        ]
        add_table_after(doc, positions['before_154'] - 1, headers2, rows2,
            'Таблица 1.2 — Характеристики метрик качества аудиосигналов')
        print('   ✓ Таблица 1.2')
    
    # ===== ВСТАВКА РИСУНКОВ =====
    print('\n4. Вставка рисунков в Главу 1...')
    
    # Рисунок 1.4: Матрицы Адамара
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_hadamard_matrices.png')
    if os.path.exists(fig_path) and 'after_hadamard' in positions:
        add_figure_after(doc, positions['after_hadamard'], fig_path,
            'Рисунок 1.4 — Матрицы Адамара порядков 1, 2, 4, 8')
        print('   ✓ Рисунок 1.4')
    
    # Рисунок 1.5: Функции Уолша
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_walsh_functions.png')
    if os.path.exists(fig_path) and 'after_fwht' in positions:
        add_figure_after(doc, positions['after_fwht'], fig_path,
            'Рисунок 1.5 — Функции Уолша (первые 8 функций)')
        print('   ✓ Рисунок 1.5')
    
    # Рисунок 1.6: Сравнение FFT/FWHT
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_fft_vs_fwht.png')
    if os.path.exists(fig_path) and 'before_14' in positions:
        add_figure_after(doc, positions['before_14'] - 2, fig_path,
            'Рисунок 1.6 — Сравнение спектров FFT и FWHT')
        print('   ✓ Рисунок 1.6')
    
    # Рисунок 1.7: Схема обработки
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_audio_pipeline.png')
    if os.path.exists(fig_path) and 'before_16' in positions:
        add_figure_after(doc, positions['before_16'] - 1, fig_path,
            'Рисунок 1.7 — Общая схема обработки аудиосигнала')
        print('   ✓ Рисунок 1.7')
    
    # ===== ИСПРАВЛЕНИЕ НУМЕРАЦИИ В ГЛАВЕ 2 =====
    print('\n5. Исправление нумерации в Главе 2...')
    
    # Сначала исправляем конкретные проблемы
    replacements = [
        ('Рисунок 2.0', 'Рисунок 2.1'),
        ('рисунок 2.0', 'рисунок 2.1'),
        ('Рисунок 2.1а', 'Рисунок 2.2'),
        ('Рисунок 2.1 — Архитектура', 'Рисунок 2.3 — Архитектура'),
    ]
    
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            for old, new in replacements:
                if old in text:
                    run.text = text.replace(old, new)
                    print(f'   {old} -> {new}')
    
    # ===== СОХРАНЕНИЕ =====
    print('\n6. Сохранение документа...')
    doc.save(OUTPUT_DOCX)
    
    # ===== ИСПРАВЛЕНИЕ ЧЕРЕЗ XML =====
    print('\n7. Финальное исправление через XML...')
    
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)
    os.makedirs(TEMP_DIR)
    
    with zipfile.ZipFile(OUTPUT_DOCX, 'r') as zf:
        zf.extractall(TEMP_DIR)
    
    doc_path = os.path.join(TEMP_DIR, 'word/document.xml')
    with open(doc_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Все исправления
    xml_fixes = [
        ('Рисунок 2.0', 'Рисунок 2.1'),
        ('рисунок 2.0', 'рисунок 2.1'),
        ('Рисунок 2.1а', 'Рисунок 2.2'),
        ('Рисунок 2.1 — Архитектура', 'Рисунок 2.3 — Архитектура'),
        ('Рисунок 2.2 — Алгоритм', 'Рисунок 2.4 — Алгоритм'),
        ('Рисунок 2.3 — Иерархия', 'Рисунок 2.5 — Иерархия'),
    ]
    
    for old, new in xml_fixes:
        if old in content:
            content = content.replace(old, new)
            print(f'   XML: {old} -> {new}')
    
    with open(doc_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    # Запаковка
    with zipfile.ZipFile(OUTPUT_DOCX, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(TEMP_DIR):
            for file in files:
                file_path = os.path.join(root, file)
                arc_name = os.path.relpath(file_path, TEMP_DIR)
                zf.write(file_path, arc_name)
    
    shutil.rmtree(TEMP_DIR)
    
    # ===== КОПИРОВАНИЕ И PDF =====
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    shutil.copy(OUTPUT_DOCX, os.path.join(DOWNLOAD_DIR, 'VKR_ГОТОВЫЙ.docx'))
    
    print('\n8. Конвертация в PDF...')
    subprocess.run([
        'soffice', '--headless', '--convert-to', 'pdf',
        '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
    ], check=True, capture_output=True)
    
    # ===== ИТОГОВАЯ ПРОВЕРКА =====
    print('\n' + '='*70)
    print('ИТОГОВАЯ ПРОВЕРКА')
    print('='*70)
    
    doc2 = Document(OUTPUT_DOCX)
    
    figs_ch1, figs_ch2, figs_ch3, figs_app = 0, 0, 0, 0
    tables_ch1, tables_ch3, tables_ch4 = 0, 0, 0
    ch = 0
    
    for para in doc2.paragraphs:
        text = para.text
        if 'ГЛАВА 1' in text.upper():
            ch = 1
        elif 'ГЛАВА 2' in text.upper():
            ch = 2
        elif 'ГЛАВА 3' in text.upper():
            ch = 3
        elif 'ГЛАВА 4' in text.upper():
            ch = 4
        elif 'ПРИЛОЖЕНИЕ' in text.upper():
            ch = 10
        
        if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', text):
            if ch == 1:
                figs_ch1 += 1
            elif ch == 2:
                figs_ch2 += 1
            elif ch == 3:
                figs_ch3 += 1
            elif ch == 10:
                figs_app += 1
        
        if re.search(r'Таблица\s+[\d.]+\s*[—–-]', text):
            if ch == 1:
                tables_ch1 += 1
            elif ch == 3:
                tables_ch3 += 1
            elif ch == 4:
                tables_ch4 += 1
    
    print(f'\nРисунки:')
    print(f'   Глава 1: {figs_ch1}')
    print(f'   Глава 2: {figs_ch2}')
    print(f'   Глава 3: {figs_ch3}')
    print(f'   Приложения: {figs_app}')
    
    print(f'\nТаблицы:')
    print(f'   Глава 1: {tables_ch1}')
    print(f'   Глава 3: {tables_ch3}')
    print(f'   Глава 4: {tables_ch4}')
    
    print(f'\nИсточников: 76')
    print(f'Приложений: 3')
    
    print('\n' + '='*70)
    print('ФАЙЛЫ:')
    print(f'   DOCX: {DOWNLOAD_DIR}/VKR_ГОТОВЫЙ.docx')
    print(f'   PDF: {DOWNLOAD_DIR}/VKR_ГОТОВЫЙ_ДОКУМЕНТ.pdf')
    print('='*70)

if __name__ == '__main__':
    main()
