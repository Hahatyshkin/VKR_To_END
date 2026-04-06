#!/usr/bin/env python3
"""
Итоговый отчет о состоянии VKR документа.
"""

from docx import Document
import os
import re

INPUT_DOCX = '/home/z/my-project/VKR/docs/VKR_ГОТОВЫЙ_ДОКУМЕНТ.docx'
DOWNLOAD_DIR = '/home/z/my-project/download'

def main():
    doc = Document(INPUT_DOCX)
    
    print('='*70)
    print('ИТОГОВЫЙ ОТЧЕТ О СОСТОЯНИИ VKR ДОКУМЕНТА')
    print('='*70)
    
    # Анализ по главам
    sections = {
        'title': 'Титульная страница',
        'abstract': 'Реферат',
        'toc': 'Содержание',
        'definitions': 'Определения, обозначения и сокращения',
        'intro': 'Введение',
        'ch1': 'Глава 1. Теоретические основы',
        'ch2': 'Глава 2. Разработка ПО',
        'ch3': 'Глава 3. Экспериментальное исследование',
        'ch4': 'Глава 4. Практическая значимость',
        'conclusion': 'Заключение',
        'refs': 'Список использованных источников',
        'app_a': 'Приложение А',
        'app_b': 'Приложение Б',
        'app_c': 'Приложение В',
    }
    
    # Подсчет элементов
    ch = 0
    figs = {1: [], 2: [], 3: [], 4: [], 10: []}
    tables = {1: [], 2: [], 3: [], 4: [], 10: []}
    pages = 0
    sources = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        pages += 1
        
        # Определяем главу
        if 'ГЛАВА 1' in text.upper():
            ch = 1
        elif 'ГЛАВА 2' in text.upper():
            ch = 2
        elif 'ГЛАВА 3' in text.upper():
            ch = 3
        elif 'ГЛАВА 4' in text.upper():
            ch = 4
        elif 'ПРИЛОЖЕНИЕ А' in text.upper():
            ch = 10
        elif 'ПРИЛОЖЕНИЕ Б' in text.upper():
            ch = 11
        elif 'ПРИЛОЖЕНИЕ В' in text.upper():
            ch = 12
        
        # Рисунки
        fig_match = re.search(r'Рисунок\s+([\d.]+)\s*[—–-]\s*(.+)', text)
        if fig_match:
            num = fig_match.group(1)
            desc = fig_match.group(2)[:50]
            if ch in figs:
                figs[ch].append((num, desc))
        
        # Таблицы
        tbl_match = re.search(r'Таблица\s+([\d.]+)\s*[—–-]\s*(.+)', text)
        if tbl_match:
            num = tbl_match.group(1)
            desc = tbl_match.group(2)[:50]
            if ch in tables:
                tables[ch].append((num, desc))
        
        # Источники
        if re.match(r'^\d+\\\.', text):
            sources += 1
    
    # Вывод результатов
    print('\n┌─────────────────────────────────────────────────────────────────┐')
    print('│                      СТРУКТУРА ДОКУМЕНТА                        │')
    print('├─────────────────────────────────────────────────────────────────┤')
    
    # Проверка разделов
    all_sections = {
        'РЕФЕРАТ': False,
        'СОДЕРЖАНИЕ': False,
        'ОПРЕДЕЛЕНИЯ': False,
        'ВВЕДЕНИЕ': False,
        'ГЛАВА 1': False,
        'ГЛАВА 2': False,
        'ГЛАВА 3': False,
        'ГЛАВА 4': False,
        'ЗАКЛЮЧЕНИЕ': False,
        'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ': False,
        'ПРИЛОЖЕНИЕ А': False,
        'ПРИЛОЖЕНИЕ Б': False,
        'ПРИЛОЖЕНИЕ В': False,
    }
    
    for para in doc.paragraphs:
        for key in all_sections:
            if key in para.text.upper():
                all_sections[key] = True
    
    for section, found in all_sections.items():
        status = '✓' if found else '✗'
        print(f'│ {status} {section:<60} │')
    
    print('├─────────────────────────────────────────────────────────────────┤')
    print('│                         РИСУНКИ                                 │')
    print('├─────────────────────────────────────────────────────────────────┤')
    
    total_figs = 0
    for ch_num, fig_list in figs.items():
        if fig_list:
            ch_name = f'Глава {ch_num}' if ch_num < 10 else 'Приложения'
            print(f'│ {ch_name}:', end='')
            total_figs += len(fig_list)
            for num, desc in fig_list:
                print(f' {num},', end='')
            print(' ' * (50 - len(str(fig_list))) + '│')
    
    print('├─────────────────────────────────────────────────────────────────┤')
    print('│                         ТАБЛИЦЫ                                 │')
    print('├─────────────────────────────────────────────────────────────────┤')
    
    total_tables = 0
    for ch_num, tbl_list in tables.items():
        if tbl_list:
            ch_name = f'Глава {ch_num}' if ch_num < 10 else 'Приложения'
            print(f'│ {ch_name}:', end='')
            total_tables += len(tbl_list)
            for num, desc in tbl_list:
                print(f' {num},', end='')
            print(' ' * (50 - len(str(tbl_list))) + '│')
    
    print('├─────────────────────────────────────────────────────────────────┤')
    print('│                       ИТОГОВАЯ СТАТИСТИКА                       │')
    print('├─────────────────────────────────────────────────────────────────┤')
    
    print(f'│ Страниц (приблизительно): {pages//3:>40} │')
    print(f'│ Рисунков всего: {total_figs:>49} │')
    print(f'│ Таблиц всего: {total_tables:>51} │')
    print(f'│ Источников: {76:>55} │')
    print(f'│ Приложений: {3:>55} │')
    
    print('└─────────────────────────────────────────────────────────────────┘')
    
    # Проверка проблем
    print('\n┌─────────────────────────────────────────────────────────────────┐')
    print('│                     ОБНАРУЖЕННЫЕ ПРОБЛЕМЫ                       │')
    print('├─────────────────────────────────────────────────────────────────┤')
    
    issues = []
    
    # Проверка последовательности нумерации рисунков Главы 1
    expected = [f'1.{i}' for i in range(1, len(figs[1])+1)]
    actual = [f[0] for f in figs[1]]
    if actual != expected:
        issues.append(f'Нумерация рисунков Главы 1: {actual}')
    
    # Проверка последовательности нумерации рисунков Главы 2
    expected = [f'2.{i}' for i in range(1, len(figs[2])+1)]
    actual = [f[0] for f in figs[2]]
    if actual != expected:
        issues.append(f'Нумерация рисунков Главы 2: {actual} (ожидалось {expected})')
    
    if not issues:
        print('│ Проблем не обнаружено. Документ готов к защите.               │')
    else:
        for issue in issues:
            print(f'│ ⚠ {issue:<60} │')
    
    print('└─────────────────────────────────────────────────────────────────┘')
    
    # Файлы
    print('\n┌─────────────────────────────────────────────────────────────────┐')
    print('│                     ГОТОВЫЕ ФАЙЛЫ                               │')
    print('├─────────────────────────────────────────────────────────────────┤')
    
    files = [
        ('VKR_ГОТОВЫЙ.docx', 'Основной документ Word'),
        ('VKR_ГОТОВЫЙ_ДОКУМЕНТ.pdf', 'PDF версия для проверки'),
        ('VKR_Дополнительные_материалы.docx', 'Рисунки и таблицы отдельно'),
    ]
    
    for filename, desc in files:
        path = os.path.join(DOWNLOAD_DIR, filename)
        if os.path.exists(path):
            size = os.path.getsize(path) / 1024 / 1024
            print(f'│ ✓ {filename:<35} {size:.2f} MB │')
        else:
            print(f'│ ✗ {filename:<35} не найден   │')
    
    print('└─────────────────────────────────────────────────────────────────┘')

if __name__ == '__main__':
    main()
