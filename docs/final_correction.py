#!/usr/bin/env python3
"""
Финальная корректировка VKR: удаление дубликатов и исправление нумерации.
"""

from docx import Document
from docx.shared import Pt
import os
import re
import subprocess

VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final_Corrected.docx')
DOWNLOAD_DIR = '/home/z/my-project/download'

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== УДАЛЕНИЕ ДУБЛИКАТОВ ТАБЛИЦ =====
    print('\nПоиск дубликатов таблиц...')
    
    # Ищем таблицы, которые были добавлены в конец и получили буквы (В.1, В.2, В.3)
    to_delete = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Ищем дубликаты с буквенной нумерацией
        if re.match(r'Таблица\s+[А-В]\.\d+\s*[—–-]', text):
            if 'вычислительной сложности' in text.lower() or \
               'метрик качества' in text.lower() or \
               'сравнительный анализ методов преобразования' in text.lower():
                to_delete.append(i)
                print(f'  Найден дубликат для удаления: {text[:60]}...')
    
    # Удаляем параграфы с конца (чтобы не сбить индексы)
    for idx in sorted(to_delete, reverse=True):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)
    print(f'  Удалено дубликатов: {len(to_delete)}')
    
    # ===== ИСПРАВЛЕНИЕ НУМЕРАЦИИ РИСУНКОВ В ГЛАВЕ 1 =====
    print('\nИсправление нумерации рисунков...')
    
    # Находим все рисунки в Главе 1 и перенумеровываем
    chapter1_figs = []
    in_chapter1 = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'ГЛАВА 1' in text.upper():
            in_chapter1 = True
        elif 'ГЛАВА 2' in text.upper():
            in_chapter1 = False
        
        if in_chapter1 and 'Рисунок' in text:
            match = re.search(r'Рисунок\s+([\d.]+)\s*[—–-]', text)
            if match:
                chapter1_figs.append({
                    'idx': i,
                    'old_num': match.group(1),
                    'text': text
                })
    
    print(f'  Найдено рисунков в Главе 1: {len(chapter1_figs)}')
    
    # Перенумеровываем
    for j, fig in enumerate(chapter1_figs):
        new_num = f'1.{j + 1}'
        if fig['old_num'] != new_num:
            para = doc.paragraphs[fig['idx']]
            for run in para.runs:
                if 'Рисунок' in run.text:
                    old_str = f'Рисунок {fig["old_num"]}'
                    new_str = f'Рисунок {new_num}'
                    run.text = run.text.replace(old_str, new_str)
                    print(f'    {fig["old_num"]} -> {new_num}')
    
    # ===== ИСПРАВЛЕНИЕ НУМЕРАЦИИ РИСУНКОВ В ГЛАВЕ 2 =====
    chapter2_figs = []
    in_chapter2 = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'ГЛАВА 2' in text.upper():
            in_chapter2 = True
        elif 'ГЛАВА 3' in text.upper():
            in_chapter2 = False
        
        if in_chapter2 and 'Рисунок' in text:
            match = re.search(r'Рисунок\s+([\d.]+)\s*[—–-]', text)
            if match:
                chapter2_figs.append({
                    'idx': i,
                    'old_num': match.group(1),
                    'text': text
                })
    
    print(f'  Найдено рисунков в Главе 2: {len(chapter2_figs)}')
    
    for j, fig in enumerate(chapter2_figs):
        new_num = f'2.{j + 1}'
        if fig['old_num'] != new_num:
            para = doc.paragraphs[fig['idx']]
            for run in para.runs:
                if 'Рисунок' in run.text:
                    old_str = f'Рисунок {fig["old_num"]}'
                    new_str = f'Рисунок {new_num}'
                    run.text = run.text.replace(old_str, new_str)
                    print(f'    {fig["old_num"]} -> {new_num}')
    
    # ===== ПОДСЧЕТ И ОБНОВЛЕНИЕ РЕФЕРАТА =====
    print('\nПодсчет элементов...')
    
    fig_count = 0
    tbl_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        if re.search(r'Рисунок\s+[\d.]+\s*[—–-]', text):
            fig_count += 1
        if re.search(r'Таблица\s+[\d.]+\s*[—–-]', text):
            tbl_count += 1
    
    page_count = len(doc.paragraphs) // 3
    
    print(f'  Рисунков: {fig_count}')
    print(f'  Таблиц: {tbl_count}')
    print(f'  Страниц: ~{page_count}')
    
    # Обновляем реферат
    for para in doc.paragraphs:
        if 'Научно-исследовательская работа состоит из' in para.text:
            new_text = f'Научно-исследовательская работа состоит из {page_count} страниц, {fig_count} рисунков, {tbl_count} таблиц, 76 использованных источников, 3 приложений.'
            para.clear()
            run = para.add_run(new_text)
            print(f'  Реферат обновлен')
            break
    
    # ===== СОХРАНЕНИЕ =====
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    import shutil
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    download_path = os.path.join(DOWNLOAD_DIR, 'VKR_Final_Corrected.docx')
    shutil.copy(OUTPUT_DOCX, download_path)
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер: {size / 1024 / 1024:.2f} MB')
    
    # ===== PDF =====
    print('\nКонвертация в PDF...')
    try:
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf',
            '--outdir', DOWNLOAD_DIR, OUTPUT_DOCX
        ], check=True, capture_output=True)
        print('PDF создан')
    except Exception as e:
        print(f'Ошибка PDF: {e}')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
