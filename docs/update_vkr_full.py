#!/usr/bin/env python3
"""
Полное обновление VKR: добавление рисунков и таблиц.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os
import re

# Пути
VKR_DIR = '/home/z/my-project/VKR/docs'
INPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Final.docx')
OUTPUT_DOCX = os.path.join(VKR_DIR, 'VKR_Complete_Updated.docx')
NEW_FIGURES_DIR = os.path.join(VKR_DIR, 'media/new_figures')

def set_cell_shading(cell, fill_color):
    """Установка цвета фона ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def insert_paragraph_after(doc, paragraph, text="", style=None):
    """Вставка параграфа после указанного."""
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

def add_table_at_position(doc, headers, rows, caption_text, insert_after_idx):
    """Добавление таблицы после указанного индекса."""
    # Получаем параграф для вставки
    para = doc.paragraphs[insert_after_idx]
    
    # Добавляем заголовок таблицы
    caption = insert_paragraph_after(doc, para, caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.bold = True
    
    # Добавляем таблицу после заголовка
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Перемещаем таблицу после заголовка
    caption._p.addnext(table._tbl)
    
    # Заголовки таблицы
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9D9D9')
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    
    # Данные таблицы
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx+1].cells[col_idx]
            cell.text = str(cell_data)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    return table

def add_figure_at_position(doc, image_path, caption_text, insert_after_idx, width_inches=5.5):
    """Добавление рисунка после указанного индекса."""
    # Получаем параграф для вставки
    para = doc.paragraphs[insert_after_idx]
    
    # Добавляем параграф с изображением
    fig_para = insert_paragraph_after(doc, para)
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    
    # Добавляем подпись
    caption_para = insert_paragraph_after(doc, fig_para)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(11)
    
    return fig_para

def main():
    print('Открытие документа...')
    doc = Document(INPUT_DOCX)
    
    # ===== АНАЛИЗ СТРУКТУРЫ =====
    print('\nАнализ структуры документа...')
    
    # Найти все заголовки главы 1
    chapter1_sections = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '1.1.' in text[:10] or '1.2.' in text[:10] or '1.3.' in text[:10] or '1.4.' in text[:10] or '1.5.' in text[:10] or '1.6.' in text[:10]:
            # Это заголовок раздела
            section_num = text.split()[0] if text.split() else ''
            chapter1_sections[section_num] = i
            print(f'  Раздел {section_num}: параграф {i}')
    
    # ===== ДОБАВЛЕНИЕ ТАБЛИЦ =====
    print('\nДобавление таблиц...')
    
    # Таблица 1.1: Сравнение вычислительной сложности
    # Вставляем после раздела 1.3.6 (перед 1.4)
    headers = ['Метод', 'Операции', 'Сложность', 'Умножения']
    rows = [
        ['FFT', 'N·log₂(N)', 'O(N log N)', 'Комплексные'],
        ['FWHT', 'N·log₂(N)', 'O(N log N)', 'Нет'],
        ['DCT', 'N·log₂(N)', 'O(N log N)', 'Вещественные'],
        ['DWT (Хаар)', 'N', 'O(N)', 'Нет'],
        ['μ-law', 'N', 'O(N)', 'Нет'],
    ]
    
    # Найти параграф перед разделом 1.4
    idx_14 = -1
    for i, para in enumerate(doc.paragraphs):
        if '1.4' in para.text and 'Интеграция' in para.text:
            idx_14 = i
            break
    
    if idx_14 > 0:
        try:
            add_table_at_position(doc, headers, rows, 
                'Таблица 1.1 — Сравнение вычислительной сложности методов преобразования',
                idx_14 - 1)
            print('  Добавлена Таблица 1.1 (вычислительная сложность)')
        except Exception as e:
            print(f'  Ошибка добавления таблицы 1.1: {e}')
    
    # Таблица 1.2: Метрики качества
    headers2 = ['Метрика', 'Область', 'Диапазон', 'Описание']
    rows2 = [
        ['SNR', 'Временная', '0-∞ дБ', 'Отношение сигнал/шум'],
        ['SI-SDR', 'Временная', '-∞-∞ дБ', 'Масштабно-инвариантное'],
        ['LSD', 'Спектральная', '0-∞ дБ', 'Спектральное расстояние'],
        ['STOI', 'Психоакуст.', '0-1', 'Разборчивость речи'],
        ['PESQ', 'Психоакуст.', '-0.5-4.5', 'Качество речи'],
    ]
    
    # Найти параграф перед 1.5.4
    idx_154 = -1
    for i, para in enumerate(doc.paragraphs):
        if '1.5.4' in para.text and 'Характеристики' in para.text:
            idx_154 = i
            break
    
    if idx_154 > 0:
        try:
            add_table_at_position(doc, headers2, rows2,
                'Таблица 1.2 — Характеристики метрик качества аудиосигналов',
                idx_154 - 1)
            print('  Добавлена Таблица 1.2 (метрики качества)')
        except Exception as e:
            print(f'  Ошибка добавления таблицы 1.2: {e}')
    
    # ===== ДОБАВЛЕНИЕ РИСУНКОВ =====
    print('\nДобавление рисунков...')
    
    # Рисунок 1.4: Матрицы Адамара
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_hadamard_matrices.png')
    if os.path.exists(fig_path):
        # Найти место после описания матриц Адамара
        for i, para in enumerate(doc.paragraphs):
            if 'H₄ = [1 1 1 1' in para.text or 'H₄ = [1 1 1 1' in para.text.replace(' ', ''):
                try:
                    add_figure_at_position(doc, fig_path,
                        'Рисунок 1.4 — Матрицы Адамара порядков 1, 2, 4, 8', i)
                    print('  Добавлен Рисунок 1.4 (матрицы Адамара)')
                except Exception as e:
                    print(f'  Ошибка добавления рисунка 1.4: {e}')
                break
    
    # Рисунок 1.5: Функции Уолша
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_walsh_functions.png')
    if os.path.exists(fig_path):
        for i, para in enumerate(doc.paragraphs):
            if 'ограниченная применимость' in para.text.lower():
                try:
                    add_figure_at_position(doc, fig_path,
                        'Рисунок 1.5 — Функции Уолша (первые 8 функций)', i)
                    print('  Добавлен Рисунок 1.5 (функции Уолша)')
                except Exception as e:
                    print(f'  Ошибка добавления рисунка 1.5: {e}')
                break
    
    # Рисунок 1.6: Сравнение спектров
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_fft_vs_fwht.png')
    if os.path.exists(fig_path):
        for i, para in enumerate(doc.paragraphs):
            if 'спектральная сходимость' in para.text.lower():
                try:
                    add_figure_at_position(doc, fig_path,
                        'Рисунок 1.6 — Сравнение спектров FFT и FWHT', i)
                    print('  Добавлен Рисунок 1.6 (сравнение спектров)')
                except Exception as e:
                    print(f'  Ошибка добавления рисунка 1.6: {e}')
                break
    
    # Рисунок 1.7: Схема обработки
    fig_path = os.path.join(NEW_FIGURES_DIR, 'fig_audio_pipeline.png')
    if os.path.exists(fig_path):
        for i, para in enumerate(doc.paragraphs):
            if '1.6' in para.text and 'Выводы' in para.text:
                try:
                    add_figure_at_position(doc, fig_path,
                        'Рисунок 1.7 — Общая схема обработки аудиосигнала', i - 1)
                    print('  Добавлен Рисунок 1.7 (схема обработки)')
                except Exception as e:
                    print(f'  Ошибка добавления рисунка 1.7: {e}')
                break
    
    # ===== ОБНОВЛЕНИЕ РЕФЕРАТА =====
    print('\nОбновление реферата...')
    
    # Подсчет итогов
    figure_count = 0
    table_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        if 'Рисунок' in text and '—' in text:
            figure_count += 1
        if 'Таблица' in text and '—' in text:
            table_count += 1
    
    print(f'  Итого рисунков: {figure_count}')
    print(f'  Итого таблиц: {table_count}')
    
    # Обновляем реферат
    for para in doc.paragraphs:
        if 'рисунков,' in para.text.lower() and 'таблиц,' in para.text.lower():
            # Нашли строку реферата с количеством
            for run in para.runs:
                if 'рисунков,' in run.text.lower():
                    # Заменяем числа
                    old_text = run.text
                    # Ищем паттерн "XX рисунков, XX таблиц"
                    new_text = re.sub(r'\d+\s+рисунков', f'{figure_count} рисунков', old_text)
                    new_text = re.sub(r'\d+\s+таблиц', f'{table_count} таблиц', new_text)
                    run.text = new_text
                    print(f'  Обновлен реферат: {old_text[:50]}... -> {new_text[:50]}...')
            break
    
    # ===== СОХРАНЕНИЕ =====
    print(f'\nСохранение в {OUTPUT_DOCX}...')
    doc.save(OUTPUT_DOCX)
    
    size = os.path.getsize(OUTPUT_DOCX)
    print(f'Размер файла: {size / 1024 / 1024:.2f} MB')
    
    print('\nГотово!')

if __name__ == '__main__':
    main()
